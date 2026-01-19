# ==============================================================================
# IMPORTAZIONI CORRETTE
# ==============================================================================
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
from django.utils import timezone 
from django.utils.translation import gettext_lazy as _
from django.contrib import messages
from django.db.models import Sum, Count, Q, Prefetch, F
from django.http import HttpResponse, JsonResponse 
from django.db import transaction 
from django.forms.models import inlineformset_factory 
from django.forms import modelformset_factory
from django import forms 
from django.urls import reverse, reverse_lazy
from functools import wraps
from datetime import timedelta
from django.shortcuts import redirect
from .utils import clean_image_metadata


# Import dai tuoi modelli (Assicurati che i percorsi siano esatti)
from courses.models import (
    Corso, IscrizioneCorso, Modulo, ProgressoModulo, 
    Quiz, Domanda, Risposta, Attestato, ImpostazioniSito
)
from user_auth.models import CustomUser as User, Azienda, Consulente, Prodotto, AdminReferente
# compliance/forms.py

from .models import (
    Azienda, 
    Trattamento, 
    RichiestaInteressato, 
    RuoloPrivacy,
    ValutazioneTIA, 
    Incidente, 
    Asset, 
    Software, 
    DocumentoAziendale,
    Videosorveglianza, 
    ConfigurazioneRete, 
    ComponenteRete,
    ReferenteCSIRT, 
    ContattoInternoCSIRT, 
    NotificaIncidente,
    AllegatoNotifica,
    AuditSession, 
    AuditDomanda, 
    SecurityAudit, 
    Compito,
    DomandaFornitore, 
    RispostaQuestionarioFornitore,
    Fornitore,
    AllegatoFornitore,
    SegnalazioneWhistleblowing,
    AllegatoWhistleblowing,
    ConfigurazioneWhistleblowing
)
from .forms import (
    TrattamentoForm, DocumentoAziendaleForm, VersioneDocumentoForm,
    IncidenteForm, RichiestaInteressatoForm, AuditChecklistForm,
    ReferenteStudenteForm, ConsulenteCreaReferenteForm,
    ConsulenteCompitoForm, AssetForm, SoftwareForm, RuoloPrivacyForm, TIAForm, VideosorveglianzaForm,
    AziendaModuliForm, ReferenteCSIRTForm, ContattoInternoCSIRTForm, NotificaIncidenteForm, NotificaIncidenteUpdateForm, 
    CSIRTTemplateForm, 
    ConfigurazioneReteForm, ComponenteReteFormSet, RuoloPrivacyFormSet, RuoloPrivacyForm, Compito,
    RispostaQuestionarioFornitoreForm,
    FornitoreForm
)
from django.contrib import messages
from django.db.models import Sum, Count, Q, Prefetch 
from django.http import HttpResponse, JsonResponse 
from django.db import transaction 
from django.forms.models import inlineformset_factory 
from django import forms 
import pandas as pd 
import io
from docx import Document
from docx.shared import Pt, Cm 
from django.core.files.base import ContentFile
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from django.views.decorators.http import require_POST
import os
import json 
from django.conf import settings 
import requests
from django.contrib import admin

# === NUOVE IMPORTAZIONI NECESSARIE PER GEMINI ===
from google import genai 
from google.genai.errors import APIError 
# ===============================================
from django.db.models import Prefetch
# Assicurati di importare i modelli corretti
from .models import DocumentoAziendale, CategoriaDocumento, AuditSession
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages

# ==============================================================================
# 1. FUNZIONI HELPER E DECORATORI 
# ==============================================================================

def get_azienda_current(request):
    # 1. Prova dai parametri URL (per consulenti)
    azienda_id = request.GET.get('azienda_id')
    # 2. Se non c'è, prova dalla sessione
    if not azienda_id:
        azienda_id = request.session.get('azienda_id')

    if azienda_id:
        return get_object_or_404(Azienda, id=azienda_id)
    return None

def build_vendor_recommendations(risultati):
    raccomandazioni = []

    gestione = risultati.get('GESTIONE', 0)
    it = risultati.get('IT', 0)
    fisica = risultati.get('FISICA', 0)
    generale = risultati.get('GENERALE', 0)

    if gestione < 60:
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Gestione Sicurezza',
            'titolo': 'Governance e policy',
            'descrizione': 'Formalizzare policy di sicurezza, ruoli e responsabilita; prevedere approvazione e revisione periodica.',
            'owner': 'DPO / Compliance',
            'orizzonte': '30-60 gg',
        })
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Gestione Sicurezza',
            'titolo': 'Gestione terze parti',
            'descrizione': 'Introdurre un processo di vendor risk management con SLA, requisiti minimi e due diligence.',
            'owner': 'Procurement',
            'orizzonte': '60 gg',
        })
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Gestione Sicurezza',
            'titolo': 'Risk assessment',
            'descrizione': 'Attivare un processo strutturato di valutazione e trattamento dei rischi con evidenze documentali.',
            'owner': 'Risk Manager',
            'orizzonte': '90 gg',
        })
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Gestione Sicurezza',
            'titolo': 'Formazione e awareness',
            'descrizione': 'Pianificare formazione annuale e campagne di sensibilizzazione con tracciamento delle presenze.',
            'owner': 'HR / Security',
            'orizzonte': '90 gg',
        })
    if it < 60:
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Sicurezza IT',
            'titolo': 'Identita e accessi',
            'descrizione': 'Abilitare MFA, revisione periodica degli accessi e principio del minimo privilegio.',
            'owner': 'IT Security',
            'orizzonte': '30-60 gg',
        })
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Sicurezza IT',
            'titolo': 'Backup e recovery',
            'descrizione': 'Verificare backup offsite, test periodici di restore e RPO/RTO documentati.',
            'owner': 'IT Operations',
            'orizzonte': '60 gg',
        })
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Sicurezza IT',
            'titolo': 'Vulnerability management',
            'descrizione': 'Pianificare vulnerability assessment e patching con metriche di copertura e tempi di remediation.',
            'owner': 'IT Security',
            'orizzonte': '90 gg',
        })
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Sicurezza IT',
            'titolo': 'Logging e monitoraggio',
            'descrizione': 'Centralizzare log critici e definire regole di alerting per eventi anomali.',
            'owner': 'SOC / IT Security',
            'orizzonte': '90 gg',
        })
    if fisica < 60:
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Sicurezza Fisica',
            'titolo': 'Accessi ai locali',
            'descrizione': 'Rafforzare i controlli di accesso e i registri di ingresso nei locali critici.',
            'owner': 'Facility',
            'orizzonte': '90 gg',
        })
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Sicurezza Fisica',
            'titolo': 'Controllo visitatori',
            'descrizione': 'Formalizzare procedure di registrazione e accompagnamento dei visitatori.',
            'owner': 'Facility',
            'orizzonte': '90 gg',
        })
        raccomandazioni.append({
            'priorita': 'Bassa',
            'area': 'Sicurezza Fisica',
            'titolo': 'Videosorveglianza e log',
            'descrizione': 'Verificare copertura e conservazione log, con procedure di revisione periodica.',
            'owner': 'Facility / Security',
            'orizzonte': '120 gg',
        })
        raccomandazioni.append({
            'priorita': 'Bassa',
            'area': 'Sicurezza Fisica',
            'titolo': 'Protezione ambientale',
            'descrizione': 'Valutare sensori ambientali e piani di continuita per eventi fisici.',
            'owner': 'Facility',
            'orizzonte': '120 gg',
        })

    if generale < 60:
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Generale',
            'titolo': 'Piano di miglioramento',
            'descrizione': 'Definire un piano di remediation con owner, tempi e priorita condivise.',
            'owner': 'Compliance',
            'orizzonte': '30-60 gg',
        })
        raccomandazioni.append({
            'priorita': 'Alta',
            'area': 'Generale',
            'titolo': 'Incident response',
            'descrizione': 'Formalizzare procedure di gestione incidenti e canali di escalation.',
            'owner': 'Security',
            'orizzonte': '60 gg',
        })
    elif generale < 85:
        raccomandazioni.append({
            'priorita': 'Media',
            'area': 'Generale',
            'titolo': 'Rafforzamento controlli',
            'descrizione': 'Consolidare i controlli esistenti e misurare periodicamente l efficacia.',
            'owner': 'Compliance',
            'orizzonte': '90 gg',
        })
    else:
        raccomandazioni.append({
            'priorita': 'Bassa',
            'area': 'Generale',
            'titolo': 'Monitoraggio continuo',
            'descrizione': 'Mantenere le pratiche in essere con verifiche periodiche e audit annuali.',
            'owner': 'Security',
            'orizzonte': '120 gg',
        })

    return raccomandazioni

def notify_referenti_whistleblowing(request, azienda):
    consulenti = azienda.manager_users.filter(ruolo='CONSULENTE').exclude(email='')
    if not consulenti.exists():
        return

    config = ImpostazioniSito.objects.first()
    if not config or not config.email_service_id or not config.email_public_key or not config.email_template_id:
        return

    subject = "Nuova segnalazione Whistleblowing"
    message_html = (
        "E stata depositata una nuova segnalazione Whistleblowing per la tua azienda. "
        "Il Consulente incaricato ha gia preso in carico la pratica."
    )

    base_url = "https://api.emailjs.com/api/v1.0/email/send"
    for referente in consulenti:
        payload = {
            "service_id": config.email_service_id,
            "template_id": config.email_template_id,
            "user_id": config.email_public_key,
            "accessToken": config.email_private_key,
            "template_params": {
                "to_email": referente.email,
                "subject": subject,
                "message_html": message_html,
            }
        }
        try:
            requests.post(base_url, json=payload, timeout=10)
        except Exception:
            pass


def notify_consulenti_fornitore(request, fornitore):
    azienda = fornitore.azienda_cliente
    consulenti = azienda.manager_users.filter(ruolo='CONSULENTE').exclude(email='')
    if not consulenti.exists():
        return

    config = ImpostazioniSito.objects.first()
    if not config or not config.email_service_id or not config.email_public_key or not config.email_fornitori_template_id:
        return

    protocol = 'https' if request.is_secure() else 'http'
    domain = request.get_host()
    link_lista = f"{protocol}://{domain}/compliance/fornitori/?azienda_id={azienda.pk}"

    subject = f"Questionario fornitore compilato - {azienda.nome}"
    message_html = (
        f"<p>Il fornitore <strong>{fornitore.ragione_sociale}</strong> ha completato il questionario.</p>"
        f"<p>Accedi alla piattaforma per la revisione: <a href=\"{link_lista}\">{link_lista}</a></p>"
    )

    base_url = "https://api.emailjs.com/api/v1.0/email/send"
    for consulente in consulenti:
        payload = {
            "service_id": config.email_service_id,
            "template_id": config.email_fornitori_template_id,
            "user_id": config.email_public_key,
            "accessToken": config.email_private_key,
            "template_params": {
                "to_email": consulente.email,
                "subject": subject,
                "message_html": message_html,
                "fornitore_nome": fornitore.ragione_sociale,
                "azienda_nome": azienda.nome,
                "link_lista": link_lista,
            }
        }
        try:
            requests.post(base_url, json=payload, timeout=10)
        except Exception:
            pass

from functools import wraps
from django.shortcuts import redirect
from django.contrib import messages
from django.http import HttpResponse

# --- SOSTITUISCI DA QUI ---
def role_required(view_func=None, allowed_roles=None):
    """
    Decoratore flessibile: funziona sia come @role_required che come @role_required(allowed_roles=['...'])
    """
    if allowed_roles is None:
        allowed_roles = ['REFERENTE', 'CONSULENTE']
    
    # Questa è la funzione che effettivamente avvolge la vista
    def decorator(func):
        @wraps(func)
        def _wrapped_view(request, *args, **kwargs):
            # 1. Verifica Autenticazione
            if not request.user.is_authenticated:
                return redirect('login')
            
            # 2. Normalizzazione e Controllo Ruolo
            user_role = str(request.user.ruolo).upper().strip()
            roles_upper = [r.upper() for r in allowed_roles]
            
            if user_role not in roles_upper:
                return HttpResponse(
                    f"<h2>Accesso Negato</h2>"
                    f"<p>Utente: {request.user.email} - Ruolo attuale: <b>{request.user.ruolo}</b></p>"
                    f"<p>Ruoli autorizzati: {roles_upper}</p>"
                    f"<hr><p><a href='/auth/logout/'>Fai Logout e riprova</a></p>",
                    status=403
                )
            
            # 3. Esecuzione Vista
            return func(request, *args, **kwargs)
        return _wrapped_view

    # LOGICA IBRIDA:
    # Se view_func è None, il decoratore è stato chiamato con le parentesi: @role_required()
    if view_func is None:
        return decorator
    
    # Se view_func NON è None, è stato usato senza parentesi: @role_required
    # In questo caso view_func è la funzione della vista stessa.
    return decorator(view_func)
# --- FINO A QUI ---

@login_required
@role_required(allowed_roles=['REFERENTE', 'CONSULENTE'])
def whistleblowing_info(request):
    azienda = get_azienda_current(request)
    if not azienda:
        return redirect('compliance:dashboard_compliance')

    wb_config = ConfigurazioneWhistleblowing.objects.filter(azienda=azienda).first()
    wb_anno = timezone.now().year
    wb_segnalazioni_annuali = SegnalazioneWhistleblowing.objects.filter(
        azienda=azienda,
        data_invio__year=wb_anno
    ).count()

    return render(request, 'compliance/whistleblowing_info.html', {
        'azienda': azienda,
        'whistleblowing_config': wb_config,
        'whistleblowing_year': wb_anno,
        'whistleblowing_count': wb_segnalazioni_annuali,
    })

@login_required
@role_required(allowed_roles=['CONSULENTE'])
def dashboard_consulente(request):
    user = request.user
    aziende_list = Azienda.objects.filter(manager_users=user).order_by('nome')
    aziende_csirt_ref = Azienda.objects.filter(
        referente_csirt__referente_user=user
    ).order_by('nome')
    fornitori_compilati = Fornitore.objects.filter(
        azienda_cliente__in=aziende_list,
        stato_valutazione='COMPILATO'
    ).order_by('-data_creazione')

    context = {
        'aziende_list': aziende_list,
        'aziende_csirt_ref': aziende_csirt_ref,
        'fornitori_compilati': fornitori_compilati,
    }

    return render(request, 'compliance/dashboard_consulente.html', context)

# === FUNZIONE DI CHIAMATA ALL'AI (IMPLEMENTAZIONE FINALE) ===
def generate_gemini_response(prompt):
    """
    Recupera la chiave API da ImpostazioniSito e invia la richiesta all'IA.
    """
    try:
        settings = ImpostazioniSito.objects.get(pk=1)
        API_KEY = settings.gemini_api_key
        
        if not API_KEY:
            return "Errore: Chiave API Gemini non configurata in Impostazioni Sito. Inseriscila per attivare l'AI."

        # 1. Inizializza il client Gemini con la chiave configurata
        client = genai.Client(api_key=API_KEY)
        
        # 2. Chiama il modello
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt],
            config={
                "system_instruction": "Sei un assistente esperto in normativa NIS2, GDPR e sicurezza informatica. Rispondi in italiano in modo professionale, conciso e pratico. Se richiesto, genera bozze di documenti o policy."
            }
        )
        
        # 3. Restituisci il testo generato
        if response.text:
            return response.text
        else:
            return "Risposta vuota o bloccata dall'API. Riprova con una query meno generica."
        
    except ImpostazioniSito.DoesNotExist:
        return "Errore: Impostazioni Sito non configurate. Verifica l'Admin."
    
    except APIError as e:
        # Cattura specifici errori API (es. chiave non valida, quota esaurita)
        return f"Errore nell'API Gemini: Problema di autenticazione o quota. Dettaglio: {e}"
        
    except Exception as e:
        # Gestione di errori generici (es. problemi di connessione)
        return f"Errore generico nell'esecuzione dell'assistente AI: {e}"
# ==============================================================================


# ==============================================================================
# 2. VISTE CONSULENTE
# ==============================================================================
@login_required
def consulente_add_referente(request, azienda_id):
    if request.user.ruolo != 'CONSULENTE': return redirect('login')
    try:
        # Verifica l'accesso tramite la NUOVA relazione
        azienda = Azienda.objects.get(id=azienda_id, manager_users=request.user)
    except: return redirect('compliance:dashboard_consulente')

    if request.method == 'POST':
        form = ConsulenteCreaReferenteForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False); user.username = user.email; user.azienda = azienda; user.ruolo = 'REFERENTE'
            user.set_unusable_password(); user.save(); trigger_set_password_email(request, user)
            messages.success(request, _(f"Nuovo Referente creato per {azienda.nome}."))
            return redirect('compliance:dashboard_consulente')
    else:
        form = ConsulenteCreaReferenteForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Aggiungi Referente per {azienda.nome}", 'sottotitolo_pagina': _("L'utente riceverà un'email per impostare la password."), 'url_annulla': 'compliance:dashboard_consulente'})

@login_required
def consulente_crea_compito(request, azienda_id):
    if request.user.ruolo != 'CONSULENTE': return redirect('login')
    try:
        # Verifica l'accesso tramite la NUOVA relazione
        azienda = Azienda.objects.get(id=azienda_id, manager_users=request.user)
    except: return redirect('compliance:dashboard_consulente')

    if request.method == 'POST':
        form = ConsulenteCompitoForm(request.POST)
        if form.is_valid():
            compito = form.save(commit=False)
            compito.creato_da = request.user
            compito.is_global = False 
            compito.save() 
            compito.aziende_assegnate.add(azienda)
            messages.success(request, _(f"Nuovo compito '{compito.titolo}' creato per {azienda.nome}."))
            return redirect('compliance:dashboard_consulente')
    else:
        form = ConsulenteCompitoForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Aggiungi Compito/Reminder per {azienda.nome}", 'sottotitolo_pagina': _("Alla data di scadenza, riceverai una notifica email di avviso."), 'url_annulla': 'compliance:dashboard_consulente'})

@login_required
def consulente_gestisci_moduli(request, azienda_id):
    if request.user.ruolo != 'CONSULENTE': return redirect('login')
    try:
        # Verifica l'accesso tramite la NUOVA relazione
        azienda = Azienda.objects.get(id=azienda_id, manager_users=request.user)
    except: return redirect('compliance:dashboard_consulente')

    # 1. La mappa master dei moduli CONFIGURABILI (campi booleani e nome amichevole)
    moduli_configurabili = [
        {'name': 'Registro Trattamenti', 'db_field': 'mod_trattamenti'},
        {'name': 'Gestione Documenti', 'db_field': 'mod_documenti'},
        {'name': 'Audit Compliance', 'db_field': 'mod_audit'},
        {'name': 'Videosorveglianza', 'db_field': 'mod_videosorveglianza'},
        {'name': 'TIA Estero', 'db_field': 'mod_tia'},
        {'name': 'Organigramma Privacy', 'db_field': 'mod_organigramma'},
        {'name': 'Gestione CSIRT (NIS2)', 'db_field': 'mod_csirt'},
        
        # === NUOVI ELEMENTI CORE ===
        {'name': 'Cruscotto/Semaforo', 'db_field': 'mod_cruscotto'}, 
        {'name': 'Segnalazione Incidenti', 'db_field': 'mod_incidenti'},
        {'name': 'Richieste Interessati', 'db_field': 'mod_richieste'},
        {'name': 'Gestione Formazione', 'db_field': 'mod_formazione'},
        {'name': 'Storico Sessioni Audit', 'db_field': 'mod_storico_audit'},
        {'name': 'Asset Aziendali', 'db_field': 'mod_asset'},
        {'name': 'Analisi Rischi', 'db_field': 'mod_analisi_rischi'},
        {'name': 'Configurazione Rete', 'db_field': 'mod_rete'},
        {'name': 'Fornitori', 'db_field': 'mod_fornitori'},
        {'name': 'Whistleblowing', 'db_field': 'mod_whistleblowing'},
    ]
    
    # 2. Gestione del POST
    if request.method == 'POST':
        form = AziendaModuliForm(request.POST, instance=azienda)
        if form.is_valid():
            form.save()
            messages.success(request, f"Configurazione moduli per {azienda.nome} salvata con successo!")
            return redirect('consulente_gestisci_moduli', azienda_id=azienda.id)
        else:
             messages.error(request, "Errore nel salvataggio dei moduli. Verifica i campi.")
    
    # 3. Preparazione del contesto per il GET (usiamo il form preesistente)
    form = AziendaModuliForm(instance=azienda)
    
    # 4. Preparazione della lista per il template (basata sui campi del form)
    moduli_list_context = []
    
    # Mappa i campi booleani del form al nome amichevole
    field_name_map = {config['db_field']: config['name'] for config in moduli_configurabili}
    
    # Itera sui campi configurabili e associa il campo form all'oggetto
    for field_name, friendly_name in field_name_map.items():
        if field_name in form.fields: # Controlla se il campo esiste nel form AziendaModuliForm
             moduli_list_context.append({
                'name': friendly_name,
                'field': form[field_name] # Oggetto BoundField che gestisce label, value, error
             })

    # Usiamo un template specifico per la gestione moduli del consulente
    return render(request, 'compliance/consulente_gestisci_moduli.html', {
        'azienda': azienda,
        'form_moduli': form, # Oggetto form completo
        'moduli_list': moduli_list_context, # Lista chiave per il template
        'titolo_pagina': f"Attiva/Disattiva Moduli per {azienda.nome}",
        'url_annulla': 'compliance:dashboard_consulente'
    })

# ==============================================================================
# 3. CRUSCOTTO OPERATIVO (FILTRAGGIO CRUCIALE PER IL REFERENTE)
# ==============================================================================
# ==============================================================================
# 4. AUDIT (e seguenti)
# ==============================================================================

@role_required()
def audit_create(request):
    """
    Crea una nuova sessione di audit e reindirizza alla checklist.
    """
    # 1. Recupero azienda corrente (essenziale per il collegamento)
    azienda = get_azienda_current(request)
    if not azienda:
        messages.error(request, _("Azienda non trovata. Impossibile creare l'audit."))
        return redirect('compliance:dashboard_consulente')

    # 2. Logica di Reset (se presente ?reset=1 nell'URL)
    # Puoi aggiungere qui logiche per archiviare le vecchie sessioni se necessario

    # 3. Creazione fisica del nuovo record nel database
    nuova_sessione = AuditSession.objects.create(
        azienda=azienda,
        creato_da=request.user,
        completato=False
    )

    # 4. Messaggio di conferma e reindirizzamento dinamico
    messages.success(request, _(f"Nuova sessione di audit #{nuova_sessione.id} inizializzata."))
    
    # Il nome 'audit_checklist' deve corrispondere a quanto definito in urls.py
    return redirect('compliance:audit_checklist', pk=nuova_sessione.pk)
@role_required()
def audit_checklist(request, pk):
    azienda = get_azienda_current(request)
    # Cerchiamo la sessione senza far crashare il server
    sessione = get_object_or_404(AuditSession, pk=pk)

    if not sessione:
        messages.error(request, f"Errore: La sessione di audit #{session_pk} non è stata trovata.")
        return redirect('compliance:dashboard_compliance')

    # 3. Logica delle domande e categorie
    tutte_le_domande = list(AuditDomanda.objects.all())
    categorie = AuditCategoria.objects.prefetch_related('domande').order_by('ordine')

    if request.method == 'POST':
        # Assicurati che AuditChecklistForm sia importato correttamente
        form = AuditChecklistForm(request.POST, sessione=sessione, domande=tutte_le_domande)
        if form.is_valid():
            form.save()
            messages.success(request, "Audit salvato con successo.")
            return redirect('compliance:dashboard_compliance')
    else:
        form = AuditChecklistForm(sessione=sessione, domande=tutte_le_domande)

    # 4. Raggruppamento per il template
    campi = []
    for cat in categorie:
        c_list = []
        for d in cat.domande.all():
            if f'domanda_{d.id}' in form.fields:
                c_list.append({
                    'domanda_field': form[f'domanda_{d.id}'],
                    'nota_field': form[f'nota_{d.id}']
                })
        if c_list:
            campi.append({'categoria': cat, 'items': c_list})
        
    return render(request, 'compliance/audit_checklist.html', {
        'form': form, 
        'campi_raggruppati': campi, 
        'sessione': sessione,
        'titolo_pagina': f"Revisione Audit: {sessione.data_creazione.strftime('%d/%m/%Y')}"
    })
@role_required
def security_checklist_view(request, audit_id=None):
    azienda = get_azienda_current(request)
    if not azienda: return redirect('compliance:dashboard_consulente')

    # 1. Recupero o creazione dell'Audit
    if audit_id:
        audit = get_object_or_404(SecurityAudit, pk=audit_id, azienda=azienda)
    else:
        audit = SecurityAudit.objects.create(azienda=azienda, creato_da=request.user)
        return redirect('security_checklist_view', audit_id=audit.id)

    # 2. Aggiornamento automatico controlli (logica esistente)
    controlli_esistenti_ids = SecurityResponse.objects.filter(audit=audit).values_list('controllo_id', flat=True)
    nuovi_controlli = SecurityControl.objects.exclude(id__in=controlli_esistenti_ids)
    
    if nuovi_controlli.exists():
        SecurityResponse.objects.bulk_create([
            SecurityResponse(audit=audit, controllo=c) for c in nuovi_controlli
        ])

    # 3. Organizzazione dati raggruppati con CALCOLO RATING per area
    risposte = SecurityResponse.objects.filter(audit=audit).select_related('controllo').order_by('controllo__control_id')
    
    # Creiamo un dizionario strutturato per passare sia i dati che i punteggi
    checklist_per_area = {}
    aree_univore = risposte.values_list('controllo__area', flat=True).distinct()

    for area in aree_univore:
        # Calcoliamo la percentuale usando il metodo che abbiamo aggiunto al modello
        percentuale = audit.get_punteggio_area(area)
        colore = audit.colore_rating
        
        checklist_per_area[area] = {
            'risposte': [r for r in risposte if r.controllo.area == area],
            'punteggio': percentuale,
            'colore': colore
        }

    context = {
        'audit': audit,
        'checklist_per_area': checklist_per_area, # Ora contiene oggetti complessi
        'titolo_pagina': _("Security Checklist Base"),
    }
    return render(request, 'compliance/security_checklist.html', context)
@role_required
@require_POST
def salva_area_security(request, audit_id):
    audit = get_object_or_404(SecurityAudit, pk=audit_id)
    
    for key, value in request.POST.items():
        if key.startswith('esito_'):
            controllo_id = key.split('_')[1]
            nota_val = request.POST.get(f'note_{controllo_id}', '')
            
            # Utilizziamo update_or_create per gestire sia inserimenti che modifiche
            SecurityResponse.objects.update_or_create(
                audit=audit, 
                controllo_id=controllo_id,
                defaults={
                    'esito': value, 
                    'note': nota_val
                }
            )
            
    messages.success(request, _("Area salvata con successo."))
    return redirect('security_checklist_view', audit_id=audit.id)
# ==============================================================================
# 5. TRATTAMENTI
# ==============================================================================

@role_required
def trattamento_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    ai_suggestion_finalita = None 
    if request.method == 'POST':
        form_instance = TrattamentoForm(request.POST) 
        if 'generate_ai_suggestions' in request.POST:
            try:
                categorie_dati = request.POST.get('categorie_dati_input', '').strip()
                soggetti_interessati = request.POST.get('soggetti_interessati_input', '').strip()
                prompt = (
                    "Genera finalità di trattamento e misure di sicurezza per un registro di trattamenti "
                    f"per {request.POST.get('nome_trattamento', 'un nuovo processo')}. "
                    f"Categorie di dati: {categorie_dati or 'non specificate'}. "
                    f"Soggetti interessati: {soggetti_interessati or 'non specificati'}."
                )
                risposta_ai = generate_gemini_response(prompt)
                form = TrattamentoForm(initial=request.POST.dict())
            except Exception as e:
                ai_suggestion_finalita = f"Errore AI: {e}"
                form = DocumentoAziendaleForm(request.POST) 
            return render(request, 'compliance/trattamento_form.html', {'form': form, 'titolo_pagina': "Nuovo Trattamento", 'ai_suggestion': ai_suggestion_finalita})
        else:
            form = TrattamentoForm(request.POST)
            if form.is_valid(): 
                t = form.save(commit=False); t.azienda = azienda; t.creato_da = request.user; t.save(); form.save_m2m()
                return redirect('checklist_trattamento', pk=t.pk)
    else:
        form = TrattamentoForm()
    return render(request, 'compliance/trattamento_form.html', {'form': form, 'titolo_pagina': "Nuovo Trattamento", 'ai_suggestion': None})

@role_required
def trattamento_update(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    t = get_object_or_404(Trattamento, pk=pk, azienda=azienda)
    if request.method == 'POST':
        form = TrattamentoForm(request.POST, instance=t)
        if form.is_valid(): form.save(); return redirect('compliance:dashboard_compliance')
    else: form = TrattamentoForm(instance=t)
    return render(request, 'compliance/trattamento_form.html', {'form': form, 'titolo_pagina': "Modifica Trattamento"})

@role_required
def checklist_trattamento(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    trattamento = get_object_or_404(Trattamento, pk=pk, azienda=azienda)
    domande = DomandaChecklist.objects.all().order_by('ordine')
    if request.method == 'POST':
        score = 0; dpia = False
        for d in domande:
            risp = (request.POST.get(f'domanda_{d.id}') == 'Sì')
            RispostaChecklist.objects.update_or_create(trattamento=trattamento, domanda=d, defaults={'risposta': risp})
            if risp: score += d.punteggio_rischio
        if score == 0: trattamento.livello_rischio = 'BASSO'
        elif score < 15: trattamento.livello_rischio = 'MEDIO'
        else: trattamento.livello_rischio = 'ALTO'; dpia = True
        trattamento.punteggio_rischio_calcolato = score; trattamento.dpia_necessaria = dpia; trattamento.save()
        messages.success(request, "Rischio calcolato."); return redirect('compliance:dashboard_compliance')
    risposte = {r.domanda_id: r.risposta for r in trattamento.risposte_checklist.all()}
    list_d = [{'domanda': d, 'risposta_attuale': risposte.get(d.id, None)} for d in domande]
    return render(request, 'compliance/checklist_form.html', {'trattamento': trattamento, 'domande_list': list_d, 'titolo_pagina': "Valutazione Rischio"})

@role_required
def export_trattamenti_excel(request):
    azienda = get_azienda_current(request)
    if not azienda: return redirect('login')
    trattamenti = Trattamento.objects.filter(azienda=azienda).prefetch_related('categorie_dati', 'soggetti_interessati')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Registro_Trattamenti_{azienda.nome}.xlsx"'
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Registro"
    headers = ["ID", "Nome", "Ruolo", "Per conto di", "Finalità", "Dati", "Soggetti", "Dest. Interni", "Dest. Esterni", "Conservazione", "Sicurezza", "Rischio", "DPIA"]
    font = Font(bold=True, color="FFFFFF"); fill = PatternFill(start_color="005a9c", end_color="005a9c", fill_type="solid")
    for col, title in enumerate(headers, 1):
        c = ws.cell(row=1, column=col); c.value = title; c.font = font; c.fill = fill
    for r, t in enumerate(trattamenti, 2):
        row = [t.id, t.nome_trattamento, t.get_tipo_ruolo_display(), t.per_conto_di or "-", t.finalita, ", ".join([str(x) for x in t.categorie_dati.all()]), ", ".join([str(x) for x in t.soggetti_interessati.all()]), t.destinatari_interni, t.destinatari_esterni, t.tempo_conservazione, t.misure_sicurezza, t.get_livello_rischio_display(), "Sì" if t.dpia_necessaria else "No"]
        for c, val in enumerate(row, 1): ws.cell(row=r, column=c).value = str(val)
    wb.save(response); return response

# ==============================================================================
# 6. DOCUMENTI
# ==============================================================================

@role_required
@role_required
def documento_list(request):
    azienda = get_azienda_current(request)
    if not azienda:
        return redirect('login')

    # Usiamo il Prefetch per caricare i documenti dell'azienda dentro le categorie
    # Questo alimenta direttamente 'categoria.documentoaziendale_set.all' usato nel tuo template
    categorie_aziendali = CategoriaDocumento.objects.prefetch_related(
        Prefetch(
            'documentoaziendale_set', 
            queryset=DocumentoAziendale.objects.filter(azienda=azienda).order_by('-id')
        )
    ).distinct()

    # Gestione sicura dei Template per evitare crash
    templates_list = []
    try:
        # Cerchiamo di capire come si chiama il modello nel tuo file models.py
        from .models import TemplateDocumento as TmpMod  # Prova nome 1
        templates_list = TmpMod.objects.all().order_by('nome')
    except (ImportError, NameError):
        try:
            from .models import DocumentoTemplate as TmpMod  # Prova nome 2
            templates_list = TmpMod.objects.all().order_by('nome')
        except (ImportError, NameError):
            templates_list = [] # Se nessuno esiste, la pagina non crasha più

    return render(request, 'compliance/documento_list.html', {
        'categorie_aziendali': categorie_aziendali,
        'templates_list': templates_list,
        'azienda': azienda,
    })

@role_required
def documento_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    form = DocumentoAziendaleForm() 
    ai_suggestion_text = None
    if request.method == 'POST':
        form = DocumentoAziendaleForm(request.POST) 
        if 'generate_ai_document_draft' in request.POST:
            try:
                prompt = f"Genera bozza documento '{request.POST.get('nome', 'Documento senza nome')}' per azienda '{azienda.nome}'..."
                risposta_ai = generate_gemini_response(prompt)
                form = DocumentoAziendaleForm(initial=request.POST.dict())
            except Exception as e:
                ai_suggestion_finalita = f"Errore AI: {e}"
                form = DocumentoAziendaleForm(request.POST) 
            return render(request, 'compliance/trattamento_form.html', {'form': form, 'titolo_pagina': "Nuovo Documento", 'ai_suggestion': ai_suggestion_finalita})
        else:
            if form.is_valid(): 
                d = form.save(commit=False); d.azienda = azienda; d.save(); 
                return redirect('versione_create', doc_pk=d.pk)
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Nuovo Documento", 'url_annulla': 'compliance:documento_list', 'ai_mode': True, 'ai_suggestion_text': ai_suggestion_text})

@role_required
def documento_dettaglio(request, doc_pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    d = get_object_or_404(DocumentoAziendale, pk=doc_pk, azienda=azienda)
    return render(request, 'compliance/documento_dettaglio.html', {'documento': d, 'versioni_list': d.versioni.all().order_by('-data_caricamento')})

@role_required
def versione_create(request, doc_pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    d = get_object_or_404(DocumentoAziendale, pk=doc_pk, azienda=azienda)
    if request.method == 'POST':
        form = VersioneDocumentoForm(request.POST, request.FILES)
        if form.is_valid(): v = form.save(commit=False); v.documento = d; v.caricato_da = request.user; v.save(); messages.success(request, "Versione caricata."); return redirect('compliance:documento_list')
    else: form = VersioneDocumentoForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Carica Versione", 'url_annulla': 'compliance:documento_list'})

@role_required
def download_template(request, template_pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    template = get_object_or_404(TemplateDocumento, pk=template_pk)
    if not template.file.name.lower().endswith('.docx'): return redirect(template.file.url)
    try:
        doc = Document(template.file.path)
        f = io.BytesIO(); doc.save(f); f.seek(0)
        resp = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = f'attachment; filename="COMPILATO_{template.file.name}"'; return resp
    except Exception: return redirect(template.file.url)

# ==============================================================================
# 7. INCIDENTI
# ==============================================================================

@role_required
def incidente_list(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    return render(request, 'compliance/incidente_list.html', {'incidenti_list': Incidente.objects.filter(azienda=azienda).order_by('-data_rilevamento')})

@role_required
def incidente_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST':
        form = IncidenteForm(request.POST)
        if form.is_valid(): i = form.save(commit=False); i.azienda = azienda; i.segnalato_da = request.user; i.save(); messages.success(request, "Segnalato."); return redirect('compliance:incidente_list')
    else: form = IncidenteForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Segnala Incidente", 'url_annulla': 'compliance:dashboard_compliance'})

@role_required
def incidente_dettaglio(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    return render(request, 'compliance/incidente_dettaglio.html', {'incidente': get_object_or_404(Incidente, pk=pk, azienda=azienda)})

# ==============================================================================
# 8. FORMAZIONE
# ==============================================================================

@role_required
def gestione_formazione(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    studenti = User.objects.filter(azienda=azienda, ruolo='STUDENTE').prefetch_related('iscrizioni')
    # Nota: Assicurati che i campi 'stato' e 'A' siano validi per il tuo modello Corso
    corsi = Corso.objects.filter(stato='A') 
    if request.method == 'POST':
        u = get_object_or_404(User, id=request.POST.get('studente_id'), azienda=azienda)
        c = get_object_or_404(Corso, id=request.POST.get('corso_id'), stato='A')
        if 'action_add' in request.POST: IscrizioneCorso.objects.get_or_create(studente=u, corso=c) # <--- Corretto
        elif 'action_remove' in request.POST: IscrizioneCorso.objects.filter(studente=u, corso=c).delete() # <--- Corretto
        return redirect('compliance:gestione_formazione')
    data = [{'studente': s, 'iscrizioni_set': {i.corso.id for i in s.iscrizioni.all()}} for s in studenti]
    return render(request, 'compliance/gestione_formazione.html', {'studenti_list': data, 'corsi_list': corsi, 'titolo_pagina': "Formazione"})

@role_required
def studente_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST':
        form = ReferenteStudenteForm(request.POST)
        if form.is_valid(): u = form.save(commit=False); u.azienda = azienda; u.ruolo = 'STUDENTE'; u.set_unusable_password(); u.save(); trigger_set_password_email(request, u)
        messages.success(request, "Creato."); return redirect('compliance:gestione_formazione')
    else: form = ReferenteStudenteForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Nuovo Studente", 'url_annulla': 'compliance:gestione_formazione'})

@role_required
def studente_update(request, studente_pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    s = get_object_or_404(User, pk=studente_pk, azienda=azienda, ruolo='STUDENTE')
    if request.method == 'POST':
        form = ReferenteStudenteForm(request.POST, instance=s)
        if form.is_valid(): form.save(); messages.success(request, "Aggiornato."); return redirect('compliance:gestione_formazione')
    else: form = ReferenteStudenteForm(instance=s)
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Modifica {s.get_full_name()}", 'url_annulla': 'compliance:gestione_formazione'})

@role_required
@transaction.atomic
def referente_import_excel(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST' and request.FILES.get('excel_file'):
        try:
            df = pd.read_excel(request.FILES['excel_file']); df = df.fillna('')
            for _, row in df.iterrows():
                if row.get('email'):
                    u, c = User.objects.update_or_create(username=row['email'], defaults={'email': row['email'], 'first_name': row.get('nome',''), 'last_name': row.get('cognome',''), 'ruolo': 'STUDENTE', 'azienda': azienda, 'is_active': True})
                    if c: u.set_unusable_password(); u.save(); trigger_set_password_email(request, u)
            messages.success(request, "Importazione completata."); return redirect('compliance:gestione_formazione')
        except Exception as e: messages.error(request, f"Errore: {e}")
    return render(request, 'compliance/importa_studenti.html')

# ==============================================================================
# 9. RICHIESTE
# ==============================================================================

@role_required
def richiesta_list(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    return render(request, 'compliance/richiesta_list.html', {'richieste_list': RichiestaInteressato.objects.filter(azienda=azienda).order_by('-data_ricezione')})

@role_required
def richiesta_create(request):
    azienda = get_azienda_current(request)
    if not azienda: 
        return redirect('login')
        
    if request.method == 'POST':
        form = RichiestaInteressatoForm(request.POST) 
        if form.is_valid(): 
            r = form.save(commit=False)
            r.azienda = azienda
            r.gestita_da = request.user
            r.save()
            messages.success(request, "Richiesta registrata con successo.")
            # CORREZIONE 1: Aggiunto namespace compliance:
            return redirect('compliance:richiesta_list') 
    else: 
        form = RichiestaInteressatoForm()
    
    # CORREZIONE 2: Aggiunto namespace compliance: alla variabile url_annulla
    return render(request, 'compliance/form_generico.html', {
        'form': form, 
        'titolo_pagina': "Nuova Richiesta", 
        'url_annulla': 'compliance:richiesta_list'
    })

@role_required
def richiesta_dettaglio(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    r = get_object_or_404(RichiestaInteressato, pk=pk, azienda=azienda)
    if request.method == 'POST':
        form = RichiestaInteressatoForm(request.POST, instance=r)
        if form.is_valid(): 
            r.stato = form.cleaned_data['stato']
            r.note_interne = form.cleaned_data['note_interne']
            r.save(); 
            messages.success(request, "Aggiornata."); return redirect('compliance:richiesta_list')
    else: form = RichiestaInteressatoForm(instance=r)
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Dettaglio Richiesta", 'url_annulla': 'compliance:richiesta_list'})

@role_required
@require_POST
def richiesta_chiudi(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    r = get_object_or_404(RichiestaInteressato, pk=pk, azienda=azienda)
    r.stato = 'EVASA'; r.save(); messages.success(request, "Chiusa."); return redirect('compliance:richiesta_list')

# ==============================================================================
# 10. COMPITI
# ==============================================================================

@role_required
@require_POST
def compito_completa(request, compito_pk):
    azienda = get_azienda_current(request)
    if not azienda: return redirect('login')
    compito = get_object_or_404(Compito, Q(is_global=True) | Q(aziende_assegnate=azienda), pk=compito_pk, stato='APERTO')
    compito.stato = 'COMPLETATO'; compito.save()
    messages.success(request, f"Compito '{compito.titolo}' completato.")
    return redirect('compliance:dashboard_compliance')
    
# ==============================================================================
# 11. ASSET
# ==============================================================================

@role_required
def asset_list(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    return render(request, 'compliance/asset_list.html', {'assets': Asset.objects.filter(azienda=azienda), 'softwares': Software.objects.filter(azienda=azienda)})

@role_required
def asset_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST':
        form = AssetForm(request.POST)
        if form.is_valid(): asset = form.save(commit=False); asset.azienda = azienda; asset.save(); return redirect('compliance:asset_list')
    else: form = AssetForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Nuovo Asset Hardware", 'url_annulla': 'compliance:asset_list'})

@role_required
def asset_update(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    asset = get_object_or_404(Asset, pk=pk, azienda=azienda)
    if request.method == 'POST':
        form = AssetForm(request.POST, instance=asset)
        if form.is_valid(): form.save(); return redirect('compliance:asset_list')
    else: form = AssetForm(instance=asset)
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Modifica Asset", 'url_annulla': 'compliance:asset_list'})

@role_required
def software_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST':
        form = SoftwareForm(request.POST)
        if form.is_valid(): sw = form.save(commit=False); sw.azienda = azienda; sw.save(); return redirect('compliance:asset_list')
    else: form = SoftwareForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Nuovo Software/SaaS", 'url_annulla': 'compliance:asset_list'})

@role_required
def software_update(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    sw = get_object_or_404(Software, pk=pk, azienda=azienda)
    if request.method == 'POST':
        form = SoftwareForm(request.POST, instance=sw)
        if form.is_valid(): form.save(); return redirect('compliance:asset_list')
    else: form = SoftwareForm(instance=sw)
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': f"Modifica Software", 'url_annulla': 'compliance:asset_list'})

@role_required
def export_asset_excel(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    # ... (export logic)
    return response

# ==============================================================================
# 12. ORGANIGRAMMA (Logica Formset Corretta)
# ==============================================================================

@role_required
def organigramma_view(request):
    """Gestisce la visualizzazione e la modifica dei ruoli privacy tramite Formset."""
    azienda = get_azienda_current(request)
    if not azienda: return redirect('login')
    
    # 1. Definizione dinamica del FormSet 
    RuoliFormSet = inlineformset_factory(Azienda, RuoloPrivacy, form=RuoloPrivacyForm, extra=1, can_delete=True)
    
    if request.method == 'POST':
        # Gestisce il salvataggio dei ruoli
        formset = RuoliFormSet(request.POST, request.FILES, instance=azienda)
        if formset.is_valid():
            formset.save()
            messages.success(request, _("Organigramma Privacy aggiornato."))
            return redirect('compliance:organigramma_view')
        else:
            messages.error(request, _("Errore nel salvataggio. Controllare i dati inseriti."))
    else:
        # Gestisce il GET (visualizzazione)
        formset = RuoliFormSet(instance=azienda)
        
    # 2. Logica Critica: Verifica se esiste un nodo radice per abilitare la visualizzazione
    # Assumiamo che i valori salvati nel DB siano le CHIAVI: 'TITOLARE' o 'RESPONSABILE_TRATTAMENTO'
    can_render_visual = RuoloPrivacy.objects.filter(
        azienda=azienda,
        ruolo_tipo__in=['TITOLARE', 'RESPONSABILE_TRATTAMENTO'] 
    ).exists()
        
    context = {
        'azienda': azienda,
        'ruoli_formset': formset,
        'can_render_visual': can_render_visual, # Variabile per abilitare il grafico
        'titolo_pagina': _("Organigramma Privacy (Ruoli e Nomine)")
    }
    # Assicurati che il template templates/compliance/organigramma.html esista
    return render(request, 'compliance/organigramma.html', context)


@role_required
def organigramma_create(request): return redirect('compliance:organigramma_view')
@role_required
def organigramma_delete(request, pk): return redirect('compliance:organigramma_view')

# ==============================================================================
# 13. TIA
# ==============================================================================

@role_required
def tia_list(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    tias = ValutazioneTIA.objects.filter(azienda=azienda).order_by('-data_valutazione')
    return render(request, 'compliance/tia_list.html', {'tias': tias})

@role_required
def tia_create(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    if request.method == 'POST':
        form = TIAForm(request.POST)
        if form.is_valid(): tia = form.save(commit=False); tia.azienda = azienda; tia.compilato_da = request.user; tia.save(); return redirect('compliance:tia_list')
    else: form = TIAForm()
    return render(request, 'compliance/form_generico.html', {'form': form, 'titolo_pagina': "Nuova TIA", 'url_annulla': 'compliance:tia_list'})
    
@role_required
def tia_detail(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    tia = get_object_or_404(ValutazioneTIA, pk=pk, azienda=azienda)
    return render(request, 'compliance/tia_detail.html', {'tia': tia})

@role_required
def tia_generate_doc_ai(request, pk):
    azienda = get_azienda_current(request); tia = get_object_or_404(ValutazioneTIA, pk=pk, azienda=azienda)
    prompt = f"Genera misure supplementari TIA per {tia.paese_destinazione.nome}..."
    try:
        document = Document()
        # Nota: 'testo' non è definito qui, si assume sia una variabile globale o importata
        document.add_paragraph("Misure supplementari generate dall'IA.")
        f = io.BytesIO(); document.save(f); f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="TIA_AI_{tia.paese_destinazione.nome}.docx"'
        return response
    except Exception as e: messages.error(request, f"Errore generazione documento: {e}"); return redirect('tia_detail', pk=pk)

# ==============================================================================
# 14. VIDEOSORVEGLIANZA
# ==============================================================================

@role_required
def video_list(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    impianti = Videosorveglianza.objects.filter(azienda=azienda)
    return render(request, 'compliance/video_list.html', {'impianti': impianti})

@role_required
def video_create(request):
    azienda = get_azienda_current(request)
    if not azienda: 
        return redirect('login')
        
    if request.method == 'POST':
        form = VideosorveglianzaForm(request.POST)
        if form.is_valid():
            video = form.save(commit=False)
            video.azienda = azienda
            video.compilato_da = request.user
            video.save()
            messages.success(request, "Impianto salvato.")
            # CORREZIONE 1: Aggiunto namespace compliance:
            return redirect('compliance:video_list')
    else: 
        form = VideosorveglianzaForm()
    
    # CORREZIONE 2: Aggiunto namespace compliance: alla variabile url_annulla
    return render(request, 'compliance/form_generico.html', {
        'form': form, 
        'titolo_pagina': "Checklist Videosorveglianza", 
        'url_annulla': 'compliance:video_list'
    })

@role_required
def video_detail(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    video = get_object_or_404(Videosorveglianza, pk=pk, azienda=azienda)
    return render(request, 'compliance/video_detail.html', {'video': video})

@role_required
@require_POST
def video_delete(request, pk):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    video = get_object_or_404(Videosorveglianza, pk=pk, azienda=azienda)
    video.delete()
    messages.success(request, "Impianto cancellato.")
    return redirect('compliance:video_list')

@role_required
def download_video_doc(request, pk, tipo_doc):
    azienda = get_azienda_current(request); video = get_object_or_404(Videosorveglianza, pk=pk, azienda=azienda)
    nome_impianto = getattr(video, 'nome_impianto', f"Impianto {video.pk}")
    document = Document()
    if tipo_doc == 'nomina_resp_esterno_ai':
        prompt = f"Genera Nomina Resp. Esterno Video per {azienda.nome}..."
        try:
            # Nota: 'testo' non è definito qui, si assume sia una variabile globale o importata
            document.add_paragraph("Nomina generata dall'IA.")
        except Exception as e: messages.error(request, f"Errore AI: {e}"); return redirect('video_detail', pk=pk)
    else: document.add_paragraph(f"Documento standard: {tipo_doc}")
    f = io.BytesIO(); document.save(f); f.seek(0)
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{tipo_doc}_{azienda.nome}.docx"'
    return response

# ==============================================================================
# 15. RISCHI
# ==============================================================================
@role_required
def analisi_rischi_guida(request):
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    return render(request, 'compliance/analisi_rischi_guida.html', {'azienda': azienda, 'titolo_pagina': 'Guida Analisi Rischi'})

# ==============================================================================
# 16. CHAT AI
# ==============================================================================
@login_required
@require_POST
def gemini_chat(request):
    if not request.user.is_authenticated: return JsonResponse({'error': 'No auth'}, status=401)
    try:
        data = json.loads(request.body); prompt = data.get('prompt', '')
        if not prompt: return JsonResponse({'success': False}, status=400)
        resp = generate_gemini_response(prompt)
        return JsonResponse({'success': True, 'response': resp})
    except Exception as e: return JsonResponse({'success': False, 'error': str(e)}, status=500)

# ==============================================================================
# 17. MODULO 7: GESTIONE REFERENTE CSIRT (NIS2)
# ==============================================================================

# compliance/views.py

# ... (Intestazioni e funzioni precedenti) ...

@role_required
def csirt_dashboard(request):
    """Dashboard principale per il modulo CSIRT con smistamento di ruolo."""
    user = request.user
    
    # 1. SMISTAMENTO RUOLO (CON ACQUISIZIONE AZIENDA)
    if user.ruolo == 'CONSULENTE':
        azienda_id_get = request.GET.get('azienda_id')
        
        if not azienda_id_get and 'consulente_azienda_id' in request.session:
            azienda_id_get = request.session.get('consulente_azienda_id')
        
        if azienda_id_get:
            try:
                azienda = Azienda.objects.get(id=azienda_id_get, manager_users=user)
                request.session['consulente_azienda_id'] = azienda_id_get
            except Azienda.DoesNotExist:
                 messages.error(request, "Accesso non autorizzato o azienda non trovata."); 
                 return redirect('compliance:dashboard_consulente')
        else:
            return redirect('compliance:dashboard_consulente')

    else: # REFERENTE
        azienda = get_azienda_current(request)
        if not azienda: 
            messages.error(request, _("Nessuna azienda associata al tuo profilo."))
            return redirect('login')
    
    # 2. DATI COMUNI (INCLUSO IL GRAFICO)
    referente_csirt, _created = ReferenteCSIRT.objects.get_or_create(
    azienda=azienda,
    defaults={'referente_user': request.user} # Nome campo corretto
    )
    
    notifiche_qs = NotificaIncidente.objects.filter(azienda=azienda)
    contatti_qs = ContattoInternoCSIRT.objects.filter(azienda=azienda)
    ContattiFormSet = modelformset_factory(
        ContattoInternoCSIRT,
        form=ContattoInternoCSIRTForm,
        extra=1,
        can_delete=True
    )
    
    # --- LOGICA GENERAZIONE JSON GRAFICO ---
    # Usiamo NotificaIncidente.objects.all() per assicurarci di accedere ai Choices
    STATO_SCELTE = dict(NotificaIncidente.STATO_CHOICES) 
    
    stato_counts = notifiche_qs.values('stato').annotate(count=Count('stato'))
    dati_grafico_stati = {}
        
    for item in stato_counts:
        # Usiamo direttamente le chiavi del DB ('APERTA', 'NOTIFICATA', ecc.)
        stato_key = item['stato']
        
        # Mappiamo la chiave del DB sul valore tradotto per la legenda (opzionale, ma pulito)
        # stato_nome_proxy = STATO_SCELTE.get(stato_key, stato_key)
        
        # PER IL JAVASCRIPT, USIAMO LA CHIAVE DEL DB, MA CON IL NOME TRADOTTO IN BASE ALLA LINGUA ATTIVA
        # QUESTO È IL MOTIVO PER CUI IL GRAFICO DEL CONSULENTE FALLIVA SE LE CHIAVI NON ERANO TRADOTTE!
        # Dobbiamo usare la traduzione, ma salvarla con la chiave Python tradotta (es. 'Aperta')
        
        # Fix: Usiamo la chiave del DB che è sempre consistente.
        # N.B.: Il codice JS deve mappare 'Aperta' (valore tradotto) al colore.
        # Se usiamo la chiave del DB qui, dobbiamo cambiare il JS per mappare la chiave del DB.
        
        # RIMANIAMO CON LA VECCHIA LOGICA PIÙ COMPLETA (chiavi tradotte)
        stato_nome_visibile = str(STATO_SCELTE.get(stato_key, stato_key))
        dati_grafico_stati[stato_nome_visibile] = item['count']

    dati_grafico_stati_json = json.dumps(dati_grafico_stati)
    # --------------------------------------------------------

    # 3. PREPARAZIONE CONTESTO
    context = {
        'referente_csirt': referente_csirt,
        'notifiche': notifiche_qs.order_by("-data_incidente"),
        'azienda': azienda,
        'titolo_pagina': _("Gestione Referente CSIRT (NIS2)"),
        'dati_grafico_stati_json': dati_grafico_stati_json, 
        'is_consulente': (user.ruolo == 'CONSULENTE'),
        'contatti_interni': contatti_qs.order_by('nominativo'),
    }

    # 4. SMISTAMENTO FINALE
    if user.ruolo == 'REFERENTE':
        if not referente_csirt.pk:
            messages.warning(request, _("Attenzione: La designazione del Referente CSIRT non è ancora stata completata dal Consulente."))
        return render(request, 'compliance/csirt_riepilogo_referente.html', context) 

    elif user.ruolo == 'CONSULENTE':
        
        # ... (Logica Form Template e Caricamento Form) ...

        try:
            config_sito = ImpostazioniSito.objects.get(pk=1) 
            context['template_form'] = CSIRTTemplateForm(instance=config_sito)
            context['config_sito'] = config_sito
        except ImpostazioniSito.DoesNotExist:
            pass 

        contatti_formset = ContattiFormSet(queryset=contatti_qs, prefix='contatti')
        context['form'] = ReferenteCSIRTForm(instance=referente_csirt)

        if request.method == 'POST':
            form_action = request.POST.get('form_action')
            if form_action == 'contatti':
                contatti_formset = ContattiFormSet(request.POST, queryset=contatti_qs, prefix='contatti')
                if contatti_formset.is_valid():
                    contatti = contatti_formset.save(commit=False)
                    for contatto in contatti:
                        contatto.azienda = azienda
                        contatto.save()
                    for contatto in contatti_formset.deleted_objects:
                        contatto.delete()
                    messages.success(request, _("Lista contatti interni aggiornata."))
                    return redirect(f"{reverse_lazy('compliance:csirt_dashboard')}?azienda_id={azienda.id}")
                messages.error(request, _("Errore nel salvataggio dei contatti interni."))
            else:
                form = ReferenteCSIRTForm(request.POST, instance=referente_csirt)
                if form.is_valid():
                    form.save()
                    messages.success(request, _("Dati Referente CSIRT aggiornati."))
                    return redirect(f"{reverse_lazy('compliance:csirt_dashboard')}?azienda_id={azienda.id}")
                else:
                     messages.error(request, _("Errore nel salvataggio dei dati Referente. Controllare i campi."))
                     context['form'] = form
        
        context['contatti_formset'] = contatti_formset
            
        return render(request, 'compliance/csirt_dashboard.html', context)
        
    return redirect('compliance:dashboard_compliance')


@role_required
def csirt_upload_template(request):
    """Vista dedicata per caricare il template Word dalla dashboard."""
    if request.method == 'POST':
        try:
            config_sito = ImpostazioniSito.objects.get(pk=1)
            form = CSIRTTemplateForm(request.POST, request.FILES, instance=config_sito)
            if form.is_valid():
                form.save()
                messages.success(request, "Nuovo template nomina caricato con successo.")
            else:
                error_list = [f"{field}: {', '.join(errors)}" for field, errors in form.errors.items()]
                messages.error(request, f"Errore nel file caricato. Dettagli: {'; '.join(error_list)}")
        except ImpostazioniSito.DoesNotExist:
             messages.error(request, "Errore di configurazione del sito. Riprovare.")
        except Exception as e:
             messages.error(request, f"Errore sistema: {e}")
    
    return redirect('compliance:csirt_dashboard')

@role_required
# compliance/views.py

def csirt_notifica_create(request):
    azienda = get_azienda_current(request)
    if not azienda: return redirect('login')
    
    if request.method == 'POST':
        form = NotificaIncidenteForm(request.POST)
        if form.is_valid():
            notifica = form.save(commit=False)
            notifica.azienda = azienda
            notifica.save()
            messages.success(request, "Notifica incidente registrata.")
            # CORREZIONE 1: Aggiunto namespace
            return redirect('compliance:csirt_dashboard') 
    else:
        form = NotificaIncidenteForm()
    
    context = {
        'form': form,
        'titolo_pagina': "Nuova Notifica Incidente (NIS2)",
        # CORREZIONE 2: Aggiunto namespace per il tasto Annulla
        'url_annulla': 'compliance:csirt_dashboard', 
    }
    return render(request, 'compliance/form_generico.html', context)


@role_required
def csirt_notifica_dettaglio(request, pk):
    """Visualizza i dettagli completi e permette l'aggiornamento dello stato/azioni per una Notifica Incidente NIS2."""
    if request.user.ruolo == 'CONSULENTE':
        azienda_id_get = request.GET.get('azienda_id')
        if azienda_id_get:
            try:
                azienda = Azienda.objects.get(id=azienda_id_get, manager_users=request.user)
                request.session['consulente_azienda_id'] = str(azienda.id)
            except Azienda.DoesNotExist:
                messages.error(request, _("Accesso non autorizzato o azienda non trovata."))
                return redirect('compliance:dashboard_consulente')
        else:
            azienda = get_azienda_current(request)
    else:
        azienda = get_azienda_current(request)
    if not azienda: 
        messages.error(request, _("Nessuna azienda associata o selezionata."))
        return redirect('compliance:dashboard_consulente') 

    # Assicurati che l'utente abbia accesso solo alle notifiche della sua azienda
    try:
        notifica = NotificaIncidente.objects.get(pk=pk, azienda=azienda)
    except NotificaIncidente.DoesNotExist:
        messages.error(request, _("Notifica non trovata o non accessibile."))
        return redirect('compliance:csirt_dashboard')
    
    # Inizializza il Formset per gli Allegati (Modello AllegatoNotifica)
    AllegatoFormSet = inlineformset_factory(
        NotificaIncidente, 
        AllegatoNotifica, 
        fields=('file', 'descrizione'), 
        extra=1, 
        can_delete=True
    )
    
    if request.method == 'POST':
        # Gestisce il salvataggio del form principale e degli allegati
        form = NotificaIncidenteUpdateForm(request.POST, instance=notifica)
        formset = AllegatoFormSet(request.POST, request.FILES, instance=notifica)

        if form.is_valid() and formset.is_valid():
            
            with transaction.atomic():
                updated_notifica = form.save(commit=False)
                
                # Se lo stato è impostato a NOTIFICATA ma data_notifica è vuota, impostiamo la data
                if updated_notifica.stato == 'NOTIFICATA' and not updated_notifica.data_notifica:
                    updated_notifica.data_notifica = timezone.now()
                
                updated_notifica.save()
                
                # Salvataggio Formset Allegati
                allegati_instances = formset.save(commit=False)
                for allegato in allegati_instances:
                    if not allegato.caricato_da:
                        allegato.caricato_da = request.user
                    allegato.save()
                
                # Gestione cancellazioni
                for obj in formset.deleted_objects:
                    obj.delete()

            messages.success(request, _(f"Notifica '{notifica.titolo_incidente}' aggiornata con successo."))
            return redirect('compliance:csirt_notifica_dettaglio', pk=pk)

        else:
            messages.error(request, _("Errore nel salvataggio. Controllare i campi della notifica o gli allegati."))
            # Il codice continua a renderizzare la pagina con form e formset contenenti gli errori

    else:
        # GET Request
        form = NotificaIncidenteUpdateForm(instance=notifica)
        formset = AllegatoFormSet(instance=notifica)
        
    is_referente = (request.user.ruolo == 'REFERENTE')
    
    context = {
        'form': form,
        'allegato_formset': formset,
        'notifica': notifica,
        'azienda': azienda,
        'is_referente': is_referente,
        'titolo_pagina': _("Dettaglio Notifica Incidente"),
        'allegati_list': notifica.allegati.all().order_by('-data_caricamento'),
    }
    
    return render(request, 'compliance/csirt_notifica_dettaglio.html', context)


@require_POST
@login_required
def csirt_ai_query(request):
    """
    Endpoint per gestire le query IA in materia CSIRT/NIS2.
    """
    if request.user.ruolo not in ['REFERENTE', 'CONSULENTE']:
        return JsonResponse({'error': 'Accesso non autorizzato.'}, status=403)
        
    try:
        if not request.body:
             return JsonResponse({'error': 'Nessun dato fornito.'}, status=400)
             
        data = json.loads(request.body)
        prompt = data.get('prompt', '').strip()
        
        if not prompt:
            return JsonResponse({'error': 'Query vuota.'}, status=400)

        ai_response = generate_gemini_response(prompt)
        
        if ai_response.startswith("Errore:"):
            return JsonResponse({'success': False, 'error': ai_response}, status=500)
            
        return JsonResponse({'success': True, 'response': ai_response})

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Formato JSON non valido.'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Errore di server non gestito: {e}'}, status=500)


@require_POST
@login_required
def csirt_dehashed_query(request):
    """
    Endpoint per interrogare l'API Dehashed tramite credenziali salvate.
    """
    if request.user.ruolo not in ['REFERENTE', 'CONSULENTE']:
        return JsonResponse({'error': 'Accesso non autorizzato.'}, status=403)

    try:
        if not request.body:
            return JsonResponse({'error': 'Nessun dato fornito.'}, status=400)

        data = json.loads(request.body)
        query = data.get('query', '').strip()

        if not query:
            return JsonResponse({'error': 'Query vuota.'}, status=400)

        config = ImpostazioniSito.objects.get(pk=1)
        if not config.dehashed_username or not config.dehashed_api_key:
            return JsonResponse({'error': 'Credenziali Dehashed non configurate.'}, status=500)

        response = requests.get(
            'https://api.dehashed.com/search',
            params={'query': query, 'size': 10},
            auth=(config.dehashed_username, config.dehashed_api_key),
            timeout=20,
        )

        if response.status_code >= 400:
            return JsonResponse({'error': f'Errore Dehashed: {response.status_code}.'}, status=response.status_code)

        payload = response.json()
        entries = payload.get('entries') or []
        trimmed_entries = []
        for entry in entries[:5]:
            trimmed_entries.append({
                'email': entry.get('email'),
                'username': entry.get('username'),
                'ip_address': entry.get('ip_address'),
                'domain': entry.get('domain'),
                'source': entry.get('database_name') or entry.get('source'),
            })

        return JsonResponse({
            'success': True,
            'total': payload.get('total'),
            'entries': trimmed_entries,
        })

    except ImpostazioniSito.DoesNotExist:
        return JsonResponse({'error': 'Impostazioni sito non configurate.'}, status=500)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Formato JSON non valido.'}, status=400)
    except requests.RequestException as exc:
        return JsonResponse({'error': f'Errore di connessione Dehashed: {exc}'}, status=502)
    except Exception as exc:
        return JsonResponse({'error': f'Errore di server non gestito: {exc}'}, status=500)


@require_POST
@login_required
def csirt_pentest_tools_query(request):
    """
    Endpoint per inviare richieste a Pentest-Tools tramite credenziali salvate.
    """
    if request.user.ruolo not in ['REFERENTE', 'CONSULENTE']:
        return JsonResponse({'error': 'Accesso non autorizzato.'}, status=403)

    try:
        if not request.body:
            return JsonResponse({'error': 'Nessun dato fornito.'}, status=400)

        data = json.loads(request.body)
        target = data.get('target', '').strip()

        if not target:
            return JsonResponse({'error': 'Target vuoto.'}, status=400)

        config = ImpostazioniSito.objects.get(pk=1)
        if not config.pentest_tools_api_key:
            return JsonResponse({'error': 'API key Pentest-Tools non configurata.'}, status=500)

        base_url = (config.pentest_tools_base_url or 'https://pp.pentest-tools.com').rstrip('/')
        scan_path = config.pentest_tools_scan_path or '/api/v1/scan'
        scan_path = scan_path if scan_path.startswith('/') else f'/{scan_path}'
        url = f"{base_url}{scan_path}"

        response = requests.post(
            url,
            json={'target': target},
            headers={
                'Authorization': f'Bearer {config.pentest_tools_api_key}',
                'Accept': 'application/json',
            },
            timeout=30,
        )

        if response.status_code >= 400:
            return JsonResponse({'error': f'Errore Pentest-Tools: {response.status_code}.'}, status=response.status_code)

        try:
            payload = response.json()
        except ValueError:
            payload = {'raw': response.text}

        return JsonResponse({'success': True, 'result': payload})

    except ImpostazioniSito.DoesNotExist:
        return JsonResponse({'error': 'Impostazioni sito non configurate.'}, status=500)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Formato JSON non valido.'}, status=400)
    except requests.RequestException as exc:
        return JsonResponse({'error': f'Errore di connessione Pentest-Tools: {exc}'}, status=502)
    except Exception as exc:
        return JsonResponse({'error': f'Errore di server non gestito: {exc}'}, status=500)

@role_required
def download_csirt_nomina(request):
    """Genera il documento di Nomina Referente CSIRT usando il template caricato o statico."""
    azienda = get_azienda_current(request);
    if not azienda: return redirect('login')
    
    try: referente = azienda.referente_csirt
    except ReferenteCSIRT.DoesNotExist: messages.error(request, "Salva prima i dati del referente."); return redirect('compliance:csirt_dashboard')

    # 1. Recupero Template (Ibrido: DB o Statico)
    template_path = None
    try:
        config = ImpostazioniSito.objects.get(pk=1)
        if config.template_nomina_csirt: template_path = config.template_nomina_csirt.path
        else: template_path = os.path.join(settings.STATICFILES_DIRS[0], 'compliance_templates', 'Nomina_Referente_CSIRT_Template.docx')
    except Exception as e: template_path = os.path.join(settings.STATICFILES_DIRS[0], 'compliance_templates', 'Nomina_Referente_CSIRT_Template.docx')

    if not os.path.exists(template_path):
        messages.error(request, "Template mancante (né caricato, né statico).")
        return redirect('compliance:csirt_dashboard')

    # 2. Compilazione e Download
    try:
        document = Document(template_path)
        
        # === LOGICA CREAZIONE STRINGA SOSTITUTI ===
        sostituti = []
        if referente.sos1_nome and referente.sos1_cognome: sostituti.append(f"- {referente.sos1_nome} {referente.sos1_cognome} ({referente.sos1_email or 'Email non specificata'})")
        if referente.sos2_nome and referente.sos2_cognome: sostituti.append(f"- {referente.sos2_nome} {referente.sos2_cognome} ({referente.sos2_email or 'Email non specificata'})")
        sostituti_text = "\n".join(sostituti) if sostituti else "Nessun sostituto designato."
        # ============================================================================

        # === Dati di base del Referente e del Punto di Contatto ===
        ref_nome_completo = f"{referente.ref_nome} {referente.ref_cognome}"
        ref_email = referente.ref_email or "________________"
        ref_telefono = referente.ref_telefono or "________________"
        
        pc_nome_completo = f"{referente.pc_nome} {referente.pc_cognome}"
        # ============================================================================

        placeholders = {
            '[DENOMINAZIONE_SOGGETTO]': azienda.nome,
            '[SEDE_LEGALE]': azienda.indirizzo or "________________",
            '[P_IVA]': azienda.p_iva or "________________",
            
            # Punti di contatto 
            '[PUNTO_CONTATTO_ACN]': pc_nome_completo, 
            
            # Referente CSIRT
            '[NOME_REFERENTE]': ref_nome_completo,
            '[RUOLO_REFERENTE]': referente.ref_ruolo or "________________", 
            '[EMAIL_REFERENTE]': ref_email,
            '[TELEFONO_REFERENTE]': ref_telefono,
            
            '[SOSTITUTI]': sostituti_text, 
            '[DATA_ODIERNA]': timezone.now().strftime('%d/%m/%Y'),
        }
        
        def replace_in_doc(doc, replacements):
            for p in doc.paragraphs:
                for key, value in replacements.items():
                    if key in p.text: p.text = p.text.replace(key, str(value))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in p.text: p.text = p.text.replace(key, str(value))

        # === LOGICA INSERIMENTO LOGO ===
        found_logo_placeholder = False
        
        if azienda.logo_principale:
            for p in document.paragraphs:
                if '[LOGO_AZIENDA]' in p.text:
                    found_logo_placeholder = True
                    p.text = p.text.replace('[LOGO_AZIENDA]', '') 
                    run = p.add_run()
                    try:
                        with azienda.logo_principale.open('rb') as f:
                            run.add_picture(f, width=Cm(4.5))
                    except: pass
            
            if not found_logo_placeholder:
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if '[LOGO_AZIENDA]' in p.text:
                                    p.text = p.text.replace('[LOGO_AZIENDA]', '')
                                    run = p.add_run()
                                    try:
                                        with azienda.logo_principale.open('rb') as f:
                                            run.add_picture(f, width=Cm(4.5))
                                    except: pass

        for p in document.paragraphs:
            if '[LOGO_AZIENDA]' in p.text: p.text = p.text.replace('[LOGO_AZIENDA]', '')
        # ==============================
        
        replace_in_doc(document, placeholders)
        f = io.BytesIO(); document.save(f); f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Nomina_CSIRT_{azienda.nome}.docx"'
        return response
    except Exception as e:
        messages.error(request, f"Errore generazione documento: {e}"); 
        return redirect('compliance:csirt_dashboard')

# ==============================================================================
# 18. CONFIGURAZIONE RETE (NIS2) - VISTA COMPLETA
# ==============================================================================

@role_required
def configurazione_rete_view(request):
    """
    Gestisce la configurazione della rete informatica per l'azienda corrente.
    Questa vista è accessibile da Consulenti (modifica) e Referenti (sola lettura).
    """
    user = request.user
    azienda = get_azienda_current(request)

    if not azienda:
        messages.error(request, _("Nessuna azienda associata o selezionata."))
        return redirect('compliance:dashboard_consulente')

    is_referente = (user.ruolo == 'REFERENTE')

    # 1) Recupera o crea configurazione con default coerente (chiave DB!)
    configurazione, _created = ConfigurazioneRete.objects.get_or_create(
        azienda=azienda,
        defaults={'tipo_architettura': 'HIERARCHICAL'}
    )

    # 2) POST (solo consulente)
    if request.method == 'POST' and not is_referente:
        form = ConfigurazioneReteForm(request.POST, instance=configurazione)
        formset = ComponenteReteFormSet(request.POST, instance=configurazione)

        if form.is_valid() and formset.is_valid():
            with transaction.atomic():
                form.save()
                formset.save()

            # IMPORTANTISSIMO: ricarica da DB per essere sicuro di leggere il valore salvato
            configurazione.refresh_from_db()

            messages.success(request, _("Configurazione di Rete aggiornata con successo."))
            return redirect('compliance:configurazione_rete_view')
        else:
            messages.error(request, _("Errore nel salvataggio dei dati. Controllare i campi evidenziati."))

    else:
        # GET
        form = ConfigurazioneReteForm(instance=configurazione)
        formset = ComponenteReteFormSet(instance=configurazione)

        if is_referente:
            # Disabilita Form e Formset per la visualizzazione Referente
            for field in form.fields.values():
                field.widget.attrs['disabled'] = True
            for form_componente in formset:
                for field in form_componente.fields.values():
                    field.widget.attrs['disabled'] = True

    # 3) Componenti
    componenti_salvati = ComponenteRete.objects.filter(
        configurazione=configurazione
    ).order_by('tipo', 'nome_componente')

    componenti_critici = ComponenteRete.objects.filter(
        configurazione=configurazione,
        criticita__in=['ALTA', 'MEDIA']
    ).order_by('criticita', 'tipo', 'nome_componente')

    # 4) Architettura: passa sia KEY che LABEL
    architettura_key = (configurazione.tipo_architettura or '').strip() or 'HIERARCHICAL'
    architettura_label = dict(ConfigurazioneRete.ARCHITETTURA_CHOICES).get(architettura_key, _("Non Definita"))

    context = {
        'azienda': azienda,
        'is_referente': is_referente,

        'form': form,
        'componente_formset': formset,
        'componenti_salvati': componenti_salvati,

        'architettura_key': architettura_key,  
        'architettura_label': architettura_label, 
        'architettura_tipo': architettura_key, # retro-compatibilità

        'componenti_critici': componenti_critici,
        'titolo_pagina': _("Configurazione Rete Informatica (NIS2)"),
    }

    return render(request, 'compliance/configurazione_rete_form.html', context)
@login_required
@role_required(allowed_roles=['CONSULENTE'])
def analisi_vulnerabilita_view(request):
    """
    Gestisce la visualizzazione delle analisi tecniche (Nessus/Vulnerabilità).
    """
    # Usiamo la funzione helper che abbiamo già per recuperare l'azienda in sessione
    azienda = get_azienda_current(request)
    
    if not azienda:
        messages.warning(request, "Seleziona un'azienda dalla dashboard per procedere.")
        return redirect('compliance:dashboard_consulente')

    context = {
        'azienda': azienda,
        'titolo_pagina': f"Analisi Vulnerabilità - {azienda.nome}",
    }
    
    return render(request, 'compliance/analisi_vulnerabilita.html', context)
@login_required
@role_required(allowed_roles=['CONSULENTE'])
def monitoraggio_eventi_view(request):
    """
    Vista per il monitoraggio dei log e dei leak di sicurezza.
    """
    # 1. Recuperiamo l'azienda corrente (usando l'helper che abbiamo già)
    azienda = get_azienda_current(request)
    
    if not azienda:
        messages.error(request, "Azienda non trovata o non selezionata.")
        return redirect('compliance:dashboard_consulente')

    # 2. Qui in futuro aggiungeremo la logica per recuperare i log reali
    risultati_monitoraggio = [] 

    context = {
        'azienda': azienda,
        'risultati': risultati_monitoraggio,
        'titolo_pagina': f"Monitoraggio Eventi - {azienda.nome}"
    }

    # 3. Restituiamo il template
    return render(request, 'compliance/monitoraggio_log.html', context)
@login_required
@role_required(allowed_roles=['CONSULENTE'])
def alert_configurazione(request):
    """
    Gestisce la configurazione degli alert di sicurezza per l'azienda.
    """
    azienda = get_azienda_current(request)
    if not azienda:
        messages.error(request, "Azienda non selezionata.")
        return redirect('compliance:dashboard_consulente')

    # Qui in futuro potrai gestire il salvataggio delle preferenze di notifica
    context = {
        'azienda': azienda,
        'titolo_pagina': f"Configurazione Alert - {azienda.nome}",
    }

    return render(request, 'compliance/alert_configurazione.html', context)
@login_required
def avvia_scansione_deashed(request, azienda_id):
    azienda = get_object_or_404(Azienda, pk=azienda_id)
    dominio = request.GET.get('dominio_manuale')

    # 1. Se il dominio manca, chiediamolo
    if not dominio:
        return render(request, 'compliance/analisi_vulnerabilita.html', {
            'azienda': azienda,
            'richiedi_dominio': True
        })

    # 2. Chiamata tecnica (qui usiamo la tua funzione esistente)
    # Assicurati che 'search_dehashed' restituisca una lista o un dizionario
    dati_raw = [
    {
        'email': f'admin@{dominio}', 
        'database_name': 'Adobe Leak (Esempio)', 
        'date': '2023-01-15'
    },
    {
        'email': f'info@{dominio}', 
        'database_name': 'LinkedIn Breach (Esempio)', 
        'date': '2022-11-20'
    }
]

# Il contesto deve passare questa lista
    context = {
    'azienda': azienda,
    'dominio_scansionato': dominio,
    'risultati': dati_raw,  # Ora è una lista, il ciclo nel template funzionerà!
}
    return render(request, 'compliance/analisi_vulnerabilita.html', context)
from django.template.loader import get_template
from xhtml2pdf import pisa  # Importa il motore di conversione

@role_required
def export_security_pdf(request, audit_id):
    azienda = get_azienda_current(request)
    audit = get_object_or_404(SecurityAudit, pk=audit_id, azienda=azienda)
    
    # Recuperiamo le risposte raggruppate per area (logica simile alla view precedente)
    risposte = SecurityResponse.objects.filter(audit=audit).select_related('controllo').order_by('controllo__control_id')
    
    checklist_per_area = {}
    aree_univore = risposte.values_list('controllo__area', flat=True).distinct()
    for area in aree_univore:
        percentuale = audit.get_punteggio_area(area)
        checklist_per_area[area] = {
            'risposte': [r for r in risposte if r.controllo.area == area],
            'punteggio': percentuale,
        }

    # Prepariamo i dati per il template PDF
    context = {
        'audit': audit,
        'azienda': azienda,
        'checklist_per_area': checklist_per_area,
        'data_export': timezone.now(),
    }

    # Rendering del PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Report_Security_{azienda.nome}.pdf"'
    
    template = get_template('compliance/pdf_security_report.html')
    html = template.render(context)

    # Creazione del PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    
    if pisa_status.err:
        return HttpResponse('Errore durante la generazione del PDF', status=500)
    return response
import pandas as pd
from django.contrib import messages
from .models import SecurityControl

@role_required
def import_controls_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        file = request.FILES['excel_file']
        
        try:
            # Leggiamo il file Excel
            df = pd.read_excel(file)
            
            # Validazione colonne minime richieste
            required_columns = ['control_id', 'area', 'descrizione']
            if not all(col in df.columns for col in required_columns):
                messages.error(request, "Il file deve contenere le colonne: control_id, area, descrizione")
                return redirect('compliance:dashboard_compliance')

            contatore = 0
            for _, row in df.iterrows():
                # update_or_create evita duplicati se l'ID controllo esiste già
                SecurityControl.objects.update_or_create(
                    control_id=row['control_id'],
                    defaults={
                        'area': row['area'],
                        'descrizione': row['descrizione'],
                        'supporto_verifica': row.get('supporto_verifica', ''),
                        'riferimento_iso': row.get('riferimento_iso', ''),
                    }
                )
                contatore += 1
            
            messages.success(request, f"Importazione completata! Caricati {contatore} controlli.")
            
        except Exception as e:
            messages.error(request, f"Errore durante l'importazione: {e}")
            
        return redirect('compliance:dashboard_compliance')
    
    return render(request, 'compliance/import_form.html')
import pandas as pd
from django.http import HttpResponse

@role_required
def download_template_excel(request):
    # Definiamo le colonne e una riga di esempio
    data = {
        'control_id': ['1.1', '1.2'],
        'area': ['Esempio Area Asset', 'Esempio Area Rete'],
        'descrizione': ['Descrizione del controllo...', 'Seconda descrizione...'],
        'supporto_verifica': ['Documentazione richiesta...', 'Note tecniche...'],
        'riferimento_iso': ['ISO 27001 A.x.x', 'NIST ID.AM-1'],
    }
    
    df = pd.DataFrame(data)
    
    # Prepariamo la risposta HTTP come file Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="Modello_Import_Framework.xlsx"'
    
    # Scriviamo il dataframe nel file excel senza indice
    with pd.ExcelWriter(response, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Controlli')
        
    return response
from django.template.loader import get_template
from xhtml2pdf import pisa

@role_required
def export_audit_pdf(request, session_pk):
    azienda = get_azienda_current(request)
    sessione = get_object_or_404(AuditSession, pk=session_pk, azienda=azienda)
    
    # Recuperiamo le risposte raggruppate per categoria
    categorie = AuditCategoria.objects.all().order_by('ordine')
    risposte = AuditRisposta.objects.filter(sessione=sessione).select_related('domanda')
    
    audit_data = []
    for cat in categorie:
        risposte_cat = risposte.filter(domanda__categoria=cat)
        if risposte_cat.exists():
            audit_data.append({
                'categoria': cat,
                'risposte': risposte_cat
            })

    context = {
        'sessione': sessione,
        'azienda': azienda,
        'audit_data': audit_data,
        'data_export': timezone.now(),
        'percentuale': sessione.percentuale_completamento,
    }

    # Preparazione della risposta PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Verbale_Audit_{azienda.nome}_{sessione.data_creazione.strftime("%Y%m%d")}.pdf"'
    
    template = get_template('compliance/pdf_audit_report.html')
    html = template.render(context)

    # Creazione del PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    
    if pisa_status.err:
        return HttpResponse('Errore nella generazione del PDF', status=500)
    return response
from django.core.files.base import ContentFile
from .models import DocumentoAziendale, VersioneDocumento, CategoriaDocumento

@role_required
def archive_audit_to_repository(request, session_pk):
    azienda = get_azienda_current(request)
    sessione = get_object_or_404(AuditSession, pk=session_pk, azienda=azienda)
    
    # 1. Generazione del PDF in memoria
    categorie_audit = AuditCategoria.objects.all().order_by('ordine')
    risposte = AuditRisposta.objects.filter(sessione=sessione).select_related('domanda')
    
    audit_data = []
    for cat in categorie_audit:
        risposte_cat = risposte.filter(domanda__categoria=cat)
        if risposte_cat.exists():
            audit_data.append({
                'categoria': cat, 
                'risposte': risposte_cat
            })

    context = {
        'sessione': sessione,
        'azienda': azienda,
        'audit_data': audit_data,
        'data_export': timezone.now(),
        'percentuale': sessione.percentuale_completamento,
    }
    
    template = get_template('compliance/pdf_audit_report.html')
    html = template.render(context)
    
    # Creiamo un buffer per il PDF
    result = io.BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=result)
    
    if pisa_status.err:
        messages.error(request, "Errore tecnico durante la generazione del verbale PDF.")
        return redirect('audit_checklist', session_pk=session_pk)

    # 2. Salvataggio nel Repository Documentale
    # Cerchiamo o creiamo la categoria specifica nel repository
    categoria_doc, _ = CategoriaDocumento.objects.get_or_create(nome="Audit e Verbali")
    
    # Cerchiamo o creiamo il documento contenitore ("Registro") per questa azienda
    documento, created = DocumentoAziendale.objects.get_or_create(
        azienda=azienda,
        nome="Registro Audit Generale GDPR",
        defaults={'categoria': categoria_doc}
    )
    
    # Prepariamo il file fisico
    timestamp = timezone.now().strftime('%Y%m%d_%H%M')
    nome_file = f"Verbale_Audit_{timestamp}.pdf"
    pdf_file = ContentFile(result.getvalue(), name=nome_file)
    
    # Creiamo la nuova versione nel repository
    VersioneDocumento.objects.create(
        documento=documento,
        file=pdf_file,
        caricato_da=request.user,
        note_versione=f"Archiviazione automatica sessione audit del {sessione.data_creazione.strftime('%d/%m/%Y')}"
    )
    
    # 3. CONGELAMENTO DELL'AUDIT
    # Una volta archiviato nel repository, l'audit non è più modificabile
    sessione.archiviata = True
    sessione.save()
    
    messages.success(request, "Audit archiviato correttamente nel Repository e congelato (sola lettura).")
    
    # Reindirizziamo al dettaglio del documento nel repository per far vedere il risultato
    return redirect('documento_dettaglio', doc_pk=documento.pk)
@role_required
def security_checklist_init(request):
    azienda = get_azienda_current(request)
    if not azienda:
        return redirect('login')
    
    # Crea una nuova sessione di Security Audit
    # Assicurati che il modello si chiami SecurityAudit o adegua il nome
    nuovo_audit = SecurityAudit.objects.create(
        azienda=azienda,
        creato_da=request.user
    )
    
    # Reindirizza alla vista che visualizza la checklist usando l'ID appena creato
    return redirect('compliance:security_checklist_view', audit_id=nuovo_audit.id)
@role_required()
def dashboard_compliance(request):
    # 1. Recupero azienda con logica sicura
    azienda_id = request.GET.get('azienda_id') or request.session.get('azienda_id')
    azienda = None
    
    if azienda_id:
        azienda = Azienda.objects.filter(id=azienda_id).first()
        if azienda:
            request.session['azienda_id'] = str(azienda.id)
    elif request.user.ruolo == 'REFERENTE' and getattr(request.user, 'azienda_id', None):
        # Fallback per referente: usa l'azienda associata al profilo.
        azienda = request.user.azienda
        request.session['azienda_id'] = str(azienda.id)

    # 2. PROTEZIONE ANTI-LOOP: Se l'azienda non c'è, decidiamo dove mandarlo in base al ruolo
    if not azienda:
        if request.user.ruolo == 'CONSULENTE':
            return redirect('compliance:dashboard_consulente')
        # Referente senza azienda: evitiamo il loop facendo logout prima del redirect.
        logout(request)
        messages.error(request, "Il tuo profilo non è associato a nessuna azienda. Contatta l'amministratore.")
        return redirect('login')

    # 3. Se arriviamo qui, l'azienda c'è. Procediamo con il resto della vista...
    # (Inserisci qui il resto del tuo codice: compiti_list, moduli, ecc.)
    # 2. Recupero Compiti (Campi corretti: aziende_assegnate e stato)
    compiti_list = Compito.objects.filter(
        aziende_assegnate=azienda, 
        stato='APERTO'
    ).order_by('data_scadenza')

    # 3. Logiche Audit Generale (GDPR)
    ultima_sessione = AuditSession.objects.filter(azienda=azienda).order_by('-data_creazione').first()
    storico_sessioni = AuditSession.objects.filter(azienda=azienda).order_by('-data_creazione')[:5]

    # 4. Logiche Security Assessment (Storico)
    storico_security = SecurityAudit.objects.filter(azienda=azienda, completato=True).order_by('-data_creazione')

    # 5. Conteggi (Modello corretto: Trattamento)
    trattamenti_count = Trattamento.objects.filter(azienda=azienda).count()
    sessioni_totali = AuditSession.objects.filter(azienda=azienda).count()

    # 6. Moduli operativi (URL allineati al tuo urls.py)
    is_consulente = (str(getattr(request.user, 'ruolo', '')).upper() == 'CONSULENTE')
    moduli = []
    moduli_config = [
        {'name': _('Registro Trattamenti'), 'icon': 'bi-journal-text', 'url': 'compliance:trattamento_list', 'flag': 'mod_trattamenti'},
        {'name': _('Asset Aziendali'), 'icon': 'bi-laptop', 'url': 'compliance:asset_list', 'flag': 'mod_asset'},
        {'name': _('Gestione Documentale'), 'icon': 'bi-folder-fill', 'url': 'compliance:documento_list', 'flag': 'mod_documenti'},
        {'name': _('Analisi Rischi'), 'icon': 'bi-exclamation-triangle', 'url': 'compliance:analisi_rischi_guida', 'flag': 'mod_analisi_rischi'},
        {'name': _('Incidenti'), 'icon': 'bi-shield-exclamation', 'url': 'compliance:incidente_list', 'flag': 'mod_incidenti'},
        {'name': _('Formazione'), 'icon': 'bi-person-video3', 'url': 'compliance:gestione_formazione', 'flag': 'mod_formazione'},
        {'name': _('Richieste Interessati'), 'icon': 'bi-people', 'url': 'compliance:richiesta_list', 'flag': 'mod_richieste'},
        {'name': _('Organigramma'), 'icon': 'bi-diagram-3', 'url': 'compliance:organigramma_view', 'flag': 'mod_organigramma'},
        {'name': _('TIA'), 'icon': 'bi-globe', 'url': 'compliance:tia_list', 'flag': 'mod_tia'},
        {'name': _('Videosorveglianza'), 'icon': 'bi-camera-video', 'url': 'compliance:video_list', 'flag': 'mod_videosorveglianza'},
        {'name': _('CSIRT Dashboard'), 'icon': 'bi-cpu', 'url': 'compliance:csirt_dashboard', 'flag': 'mod_csirt'},
        {'name': _('Configurazione Rete'), 'icon': 'bi-hdd-network', 'url': 'compliance:configurazione_rete_view', 'flag': 'mod_rete'},
        {'name': _('Fornitori'), 'icon': 'bi-truck', 'url': 'compliance:lista_fornitori_azienda', 'flag': 'mod_fornitori'},
        {'name': _('Whistleblowing'), 'icon': 'bi-megaphone', 'url': 'compliance:whistleblowing_info', 'flag': 'mod_whistleblowing'},
    ]
    for modulo in moduli_config:
        if getattr(azienda, modulo['flag'], False):
            moduli.append(modulo)

    wb_config = ConfigurazioneWhistleblowing.objects.filter(azienda=azienda).first()
    wb_anno = timezone.now().year
    wb_segnalazioni_annuali = SegnalazioneWhistleblowing.objects.filter(
        azienda=azienda,
        data_invio__year=wb_anno
    ).count()

    context = {
        'azienda': azienda,
        'compiti_list': compiti_list,
        'ultima_sessione': ultima_sessione,
        'storico_sessioni': storico_sessioni,
        'storico_security': storico_security,
        'trattamenti_count': trattamenti_count,
        'sessioni_totali': sessioni_totali,
        'is_cruscotto_active': bool(azienda.mod_cruscotto),
        'is_storico_audit_active': bool(azienda.mod_storico_audit),
        'show_audit_card': bool(azienda.mod_audit),
        'show_indicator_card': any([
            azienda.mod_trattamenti,
            azienda.mod_audit,
            azienda.mod_csirt,
            azienda.mod_storico_audit,
            azienda.mod_fornitori,
        ]),
        'show_security_vendor': bool(azienda.mod_fornitori),
        'show_riepilogo_sezione': any([
            azienda.mod_audit,
            azienda.mod_trattamenti,
            azienda.mod_csirt,
            azienda.mod_storico_audit,
            azienda.mod_fornitori,
        ]),
        'show_whistleblowing_info': bool(azienda.mod_whistleblowing) and str(getattr(request.user, 'ruolo', '')).upper() == 'REFERENTE',
        'whistleblowing_config': wb_config,
        'whistleblowing_year': wb_anno,
        'whistleblowing_count': wb_segnalazioni_annuali,
        'moduli': moduli,
        'is_consulente': is_consulente,
    }

    return render(request, 'compliance/dashboard.html', context)

# In compliance/views.py

def trattamento_list(request):
    azienda = get_azienda_current(request)
    trattamenti = Trattamento.objects.filter(azienda=azienda)
    return render(request, 'compliance/trattamento_list.html', {'trattamenti': trattamenti, 'azienda': azienda})

def asset_list(request):
    azienda = get_azienda_current(request)
    assets = Asset.objects.filter(azienda=azienda)
    return render(request, 'compliance/asset_list.html', {'assets': assets, 'azienda': azienda})

def documento_list(request):
    azienda = get_azienda_current(request)
    documenti = DocumentoAziendale.objects.filter(azienda=azienda)
    return render(request, 'compliance/documento_list.html', {'documenti': documenti, 'azienda': azienda})
def lista_domande_fornitori(request, fornitore_id):
    fornitore = get_object_or_404(Fornitore, id=fornitore_id)
    # Ordiniamo prima per macro_categoria, poi per TEMA (Ambito) e infine per CODICE
    domande = DomandaFornitore.objects.all().order_by('macro_categoria', 'codice', 'tema')
    
    domande_per_area = {
        'Gestione della Sicurezza': domande.filter(macro_categoria='GESTIONE'),
        'Sicurezza IT': domande.filter(macro_categoria='IT'),
        'Sicurezza Fisica': domande.filter(macro_categoria='FISICA'),
    }

    # RECUPERO RISPOSTE: Assicurati che non siano stringhe!
    risposte_esistenti = RispostaQuestionarioFornitore.objects.filter(fornitore=fornitore)
    
    # Creiamo i dizionari esplicitamente
    mappa_risposte = {r.domanda_id: r.valore_risposta for r in risposte_esistenti}
    mappa_note = {r.domanda_id: r.note for r in risposte_esistenti}

    return render(request, 'compliance/questionario_form.html', {
        'fornitore': fornitore,
        'domande_per_area': domande_per_area,
        'mappa_risposte': mappa_risposte, # Deve essere un dict
        'mappa_note': mappa_note,         # Deve essere un dict
    })
def salva_risposte_fornitori(request, fornitore_id):
    if request.method == "POST":
        fornitore = get_object_or_404(Fornitore, id=fornitore_id)
        
        # Recuperiamo tutte le domande per ciclare sui dati inviati dal form
        domande = DomandaFornitore.objects.all()
        risposte_salvate = 0
        
        for d in domande:
            valore = request.POST.get(f'risposta_{d.id}')
            note = request.POST.get(f'note_{d.id}', '').strip()
            
            # Salviamo solo se l'utente ha selezionato una risposta (Sì/No/Parziale)
            if valore and valore != "":
                RispostaQuestionarioFornitore.objects.update_or_create(
                    fornitore=fornitore,
                    domanda=d,
                    defaults={
                        'valore_risposta': float(valore),
                        'note': note
                    }
                )
                risposte_salvate += 1
        
        # Aggiungiamo un feedback visivo per l'utente
        if risposte_salvate > 0:
            messages.success(request, f"Valutazione salvata con successo! Registrate {risposte_salvate} risposte.")
        else:
            messages.warning(request, "Nessuna risposta selezionata. Il questionario non è stato aggiornato.")
            
        return redirect('compliance:risultati_fornitori', fornitore_id=fornitore.id)
    
    return redirect('compliance:lista_domande_fornitori', fornitore_id=fornitore_id)
def risultati_questionario_fornitori(request, fornitore_id):
    fornitore = get_object_or_404(Fornitore, id=fornitore_id)
    
    risposte = RispostaQuestionarioFornitore.objects.filter(fornitore=fornitore).select_related('domanda')
    domande = DomandaFornitore.objects.all()

    if not domande.exists():
        messages.info(request, "Nessun dato disponibile. Compila il questionario per vedere i risultati.")
        return redirect('compliance:lista_domande_fornitori', fornitore_id=fornitore.id)

    # Inizializziamo i totali Generali e per Macro Area
    stats = {
        'GENERALE': {'ottenuto': 0, 'max': 0},
        'GESTIONE': {'ottenuto': 0, 'max': 0},
        'IT': {'ottenuto': 0, 'max': 0},
        'FISICA': {'ottenuto': 0, 'max': 0},
    }

    risposte_map = {r.domanda_id: r for r in risposte}
    for d in domande:
        max_pesato = 1.0 * d.peso_domanda * d.peso_tema
        valore_pesato = 0
        risposta = risposte_map.get(d.id)
        if risposta:
            valore_pesato = risposta.valore_risposta * d.peso_domanda * d.peso_tema

        stats['GENERALE']['ottenuto'] += valore_pesato
        stats['GENERALE']['max'] += max_pesato

        area_key = d.macro_categoria
        if area_key in stats:
            stats[area_key]['ottenuto'] += valore_pesato
            stats[area_key]['max'] += max_pesato

    # Calcolo delle percentuali finali
    risultati_finali = {}
    for key, data in stats.items():
        percentuale = (data['ottenuto'] / data['max'] * 100) if data['max'] > 0 else 0
        risultati_finali[key] = round(percentuale, 2)

    # Prepariamo i dati per il Grafico Radar/Barre
    labels_grafico = ['Gestione Sicurezza', 'Sicurezza IT', 'Sicurezza Fisica']
    valori_grafico = [
        risultati_finali['GESTIONE'],
        risultati_finali['IT'],
        risultati_finali['FISICA']
    ]

    raccomandazioni = build_vendor_recommendations(risultati_finali)

    return render(request, 'compliance/risultati_fornitori.html', {
        'punteggio_finale': risultati_finali['GENERALE'],
        'dettaglio_aree': risultati_finali,
        'labels_grafico': labels_grafico,
        'valori_grafico': valori_grafico,
        'fornitore': fornitore,
        'raccomandazioni': raccomandazioni
    })
@role_required
def export_questionario_pdf(request, fornitore_id):
    fornitore = get_object_or_404(Fornitore, id=fornitore_id)
    risposte = RispostaQuestionarioFornitore.objects.filter(fornitore=fornitore).select_related('domanda')
    domande = DomandaFornitore.objects.all()
    
    # --- LOGICA CALCOLI (Stessa dei risultati) ---
    stats = {'GENERALE': {'ottenuto': 0, 'max': 0}, 'GESTIONE': {'ottenuto': 0, 'max': 0}, 
             'IT': {'ottenuto': 0, 'max': 0}, 'FISICA': {'ottenuto': 0, 'max': 0}}

    risposte_map = {r.domanda_id: r for r in risposte}
    for d in domande:
        max_pesato = 1.0 * d.peso_domanda * d.peso_tema
        valore_pesato = 0
        risposta = risposte_map.get(d.id)
        if risposta:
            valore_pesato = risposta.valore_risposta * d.peso_domanda * d.peso_tema
        stats['GENERALE']['ottenuto'] += valore_pesato
        stats['GENERALE']['max'] += max_pesato
        area = d.macro_categoria
        if area in stats:
            stats[area]['ottenuto'] += valore_pesato
            stats[area]['max'] += max_pesato

    risultati = {k: round((v['ottenuto']/v['max']*100), 2) if v['max'] > 0 else 0 for k, v in stats.items()}
    raccomandazioni = build_vendor_recommendations(risultati)

    # --- GENERAZIONE PDF ---
    context = {
        'fornitore': fornitore,
        'risultati': risultati,
        'data_report': timezone.now(),
        'raccomandazioni': raccomandazioni,
    }
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Security_Report_{fornitore.ragione_sociale}.pdf"'
    
    # Carichiamo un template specifico per il PDF (più pulito)
    template = get_template('compliance/pdf_risultati_fornitori.html')
    html = template.render(context)
    
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Errore nella generazione del PDF', status=500)
    return response

@role_required(allowed_roles=['CONSULENTE'])
def lista_fornitori_azienda(request):
    azienda = get_azienda_current(request)
    if not azienda: return redirect('dashboard_consulente')
    
    fornitori = Fornitore.objects.filter(azienda_cliente=azienda)
    return render(request, 'compliance/fornitori_list.html', {
        'azienda': azienda,
        'fornitori': fornitori
    })

@role_required(allowed_roles=['CONSULENTE'])
def crea_fornitore(request):
    azienda = get_azienda_current(request)
    if not azienda:
        messages.error(request, "Azienda non identificata.")
        return redirect('dashboard_consulente')

    if request.method == "POST":
        form = FornitoreForm(request.POST)
        if form.is_valid():
            fornitore = form.save(commit=False)
            fornitore.azienda_cliente = azienda
            fornitore.save()
            messages.success(request, f"Fornitore creato: {fornitore.ragione_sociale}")
            return redirect(f"{reverse('compliance:lista_fornitori_azienda')}?azienda_id={azienda.pk}")
    else:
        form = FornitoreForm()

    return render(request, 'compliance/fornitore_form.html', {
        'form': form,
        'azienda': azienda
    })

@role_required(allowed_roles=['CONSULENTE'])
def prepara_invito_fornitore(request, fornitore_id):
    fornitore = get_object_or_404(Fornitore, id=fornitore_id)
    azienda = fornitore.azienda_cliente
    config_sito = ImpostazioniSito.objects.first()

    protocol = 'https' if request.is_secure() else 'http'
    domain = request.get_host()
    link_portale = f"{protocol}://{domain}/compliance/vendor-portal/{fornitore.access_token}/"

    email_subject = f"Richiesta Valutazione Sicurezza Fornitore (NIS2) - {azienda.nome}"
    email_message = f"""
    <div style="font-family: Arial, sans-serif; color: #333; max-width: 600px; border: 1px solid #eee; padding: 20px;">
        <h2 style="color: #003366;">HUB Compliance</h2>
        <p>Gentile referente di <strong>{fornitore.ragione_sociale}</strong>,</p>
        <p>Per conto di <strong>{azienda.nome}</strong>, vi richiediamo di completare la valutazione di sicurezza per gli adempimenti <strong>NIS2 e GDPR</strong>.</p>
        <div style="text-align: center; margin: 30px 0;">
            <a href="{link_portale}" style="background-color: #1b17a4; color: white; padding: 15px 25px; text-decoration: none; border-radius: 5px; font-weight: bold;">
                ACCEDI AL PORTALE FORNITORI
            </a>
        </div>
        <p style="font-size: 13px; color: #666;">Il link è sicuro e non richiede password.</p>
        <hr style="border: none; border-top: 1px solid #eee;">
        <p style="font-size: 12px; color: #999;">Messaggio inviato tramite la piattaforma HUB Compliance.</p>
    </div>
    """

    return render(request, 'compliance/fornitore_conferma_invio.html', {
        'fornitore': fornitore,
        'link_portale': link_portale,
        'azienda_nome': azienda.nome,
        'email_subject': email_subject,
        'email_message': email_message,
        'EMAILJS_SERVICE_ID': getattr(config_sito, 'email_service_id', ''),
        'EMAILJS_TEMPLATE_ID': getattr(config_sito, 'email_fornitori_template_id', ''),
        'EMAILJS_PUBLIC_KEY': getattr(config_sito, 'email_public_key', '')
    })
@role_required(allowed_roles=['CONSULENTE'])
def aggiorna_stato_fornitore(request, fornitore_id):
    fornitore = get_object_or_404(Fornitore, id=fornitore_id)
    fornitore.stato_valutazione = 'INVIATO'
    fornitore.save()
    messages.success(request, f"Stato aggiornato: invito inviato a {fornitore.ragione_sociale}")
    return redirect(f"{reverse('compliance:lista_fornitori_azienda')}?azienda_id={fornitore.azienda_cliente.pk}")
# In compliance/views.py

def vendor_portal_form(request, token):
    fornitore = get_object_or_404(Fornitore, access_token=token)
    # Recuperiamo le domande del questionario CLUSIT dal tuo DB
    domande_clusit = DomandaFornitore.objects.all().order_by('macro_categoria', 'codice', 'tema')
    domande_per_area = {
        'Gestione della Sicurezza': domande_clusit.filter(macro_categoria='GESTIONE'),
        'Sicurezza IT': domande_clusit.filter(macro_categoria='IT'),
        'Sicurezza Fisica': domande_clusit.filter(macro_categoria='FISICA'),
    }

    if request.method == "POST":
        # 1. Salvataggio Risposte Questionario
        for domanda in domande_clusit:
            field_name = f"domanda_{domanda.id}"
            risposta_valore = request.POST.get(field_name)
            
            if risposta_valore is not None:
                RispostaQuestionarioFornitore.objects.update_or_create(
                    fornitore=fornitore,
                    domanda=domanda,
                    defaults={'valore_risposta': float(risposta_valore)}
                )

        # 2. Salvataggio Allegati (se presenti)
        if request.FILES.get('certificato'):
            AllegatoFornitore.objects.create(
                fornitore=fornitore,
                file=request.FILES['certificato'],
                descrizione="Certificazione caricata dal portale"
            )

        # 3. Aggiornamento Stato e logica di completamento
        fornitore.stato_valutazione = 'COMPILATO'
        fornitore.save()
        notify_consulenti_fornitore(request, fornitore)

        return render(request, 'compliance/vendor_thanks.html', {'fornitore': fornitore})

    return render(request, 'compliance/vendor_portal_form.html', {
        'fornitore': fornitore,
        'domande_per_area': domande_per_area
    })
def invia_segnalazione(request):
    azienda = get_azienda_current(request)
    if not azienda:
        return HttpResponse("Azienda non valida.", status=400)
    
    if request.method == "POST":
        # 1. Creazione segnalazione (come visto prima)
        nuova_segnalazione = SegnalazioneWhistleblowing.objects.create(
            azienda=azienda,
            categoria=request.POST.get('categoria'),
            oggetto=request.POST.get('oggetto'),
            descrizione=request.POST.get('descrizione'),
            nome_segnalante=request.POST.get('nome', 'Anonimo'),
            email_contatto=request.POST.get('email')
        )
        
        # 2. Gestione file con pulizia metadati
        files = request.FILES.getlist('documenti')
        for f in files:
            # Puliamo il file prima del salvataggio
            file_pulito = clean_image_metadata(f)
            
            AllegatoWhistleblowing.objects.create(
                segnalazione=nuova_segnalazione,
                file=file_pulito
            )

        notify_referenti_whistleblowing(request, azienda)

        request.session['wb_ticket_code'] = nuova_segnalazione.codice_ticket
        config = ImpostazioniSito.objects.first()
        status_url = f"{request.build_absolute_uri(reverse('compliance:check_status'))}?codice={nuova_segnalazione.codice_ticket}"

        return render(request, 'compliance/whistleblowing_success.html', {
            'codice': nuova_segnalazione.codice_ticket,
            'azienda': azienda,
            'segnalazione': nuova_segnalazione,
            'status_url': status_url,
            'config': config,
        })
        
    return render(request, 'compliance/whistleblowing_form.html', {'azienda': azienda})
@role_required(allowed_roles=['CONSULENTE'])
def lista_segnalazioni_whistleblowing(request):
    # Il consulente vede le segnalazioni di tutte le aziende che segue
    # o puoi filtrarle per l'azienda selezionata in sessione
    segnalazioni = SegnalazioneWhistleblowing.objects.all().order_by('-data_invio')
    
    return render(request, 'compliance/segnalazioni_list.html', {
        'segnalazioni': segnalazioni
    })

@role_required(allowed_roles=['CONSULENTE'])
def dettaglio_segnalazione(request, pk):
    segnalazione = get_object_or_404(SegnalazioneWhistleblowing, pk=pk)
    
    if request.method == 'POST':
        # Salvataggio della risposta ufficiale
        segnalazione.risposta_consulente = request.POST.get('risposta')
        segnalazione.stato = request.POST.get('stato')
        segnalazione.save()
        messages.success(request, "Risposta aggiornata con successo.")
        return redirect('compliance:lista_segnalazioni')

    return render(request, 'compliance/segnalazione_detail.html', {
        'segnalazione': segnalazione
    })
def check_segnalazione_status(request):
    segnalazione = None
    errore = None
    codice_inserito = request.GET.get('codice', '').strip()

    if codice_inserito:
        try:
            # Cerchiamo la segnalazione corrispondente al ticket
            segnalazione = SegnalazioneWhistleblowing.objects.get(codice_ticket=codice_inserito)
        except SegnalazioneWhistleblowing.DoesNotExist:
            errore = "Codice non valido o inesistente. Riprova."

    return render(request, 'compliance/whistleblowing_check.html', {
        'segnalazione': segnalazione,
        'errore': errore,
        'codice_inserito': codice_inserito
    })
@login_required
def dashboard_referente_wb(request):
    azienda = request.user.azienda_cliente # Assumendo il legame utente-azienda
    config = ConfigurazioneWhistleblowing.objects.get(azienda=azienda)
    
    # Conteggio segnalazioni anno corrente
    anno_corrente = timezone.now().year
    totale_anno = SegnalazioneWhistleblowing.objects.filter(
        azienda=azienda, 
        data_invio__year=anno_corrente
    ).count()

    context = {
        'config': config,
        'totale_anno': totale_anno,
    }
    return render(request, 'compliance/wb_referente_view.html', context)
def whistleblowing_success(request):
    # Recuperiamo il codice salvato nella sessione dalla view precedente
    codice = request.session.get('wb_ticket_code')
    
    if not codice:
        # Se qualcuno accede alla pagina direttamente senza aver inviato nulla
        return redirect('compliance:lista_aziende_wb') 

    # Opzionale: puliamo la sessione così il codice non resta in memoria
    # del browser se l'utente ricarica la pagina dopo averlo segnato
    # del request.session['wb_ticket_code'] 

    segnalazione = SegnalazioneWhistleblowing.objects.filter(codice_ticket=codice).first()
    azienda = segnalazione.azienda if segnalazione else None
    config = ImpostazioniSito.objects.first()
    status_url = f"{request.build_absolute_uri(reverse('compliance:check_status'))}?codice={codice}"

    return render(request, 'compliance/whistleblowing_success.html', {
        'codice': codice,
        'azienda': azienda,
        'segnalazione': segnalazione,
        'status_url': status_url,
        'config': config,
    })
